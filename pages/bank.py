# -*- coding: utf-8 -*-
import streamlit as st
from docx import Document
# THÊM IMPORT ĐỂ XỬ LÝ ĐỊNH DẠNG (HIGHLIGHT)
from docx.enum.text import WD_COLOR_INDEX 
import re
import math
import pandas as pd
import base64
import os
import random 
from deep_translator import GoogleTranslator

st.markdown(
    """
    <style>
    /* 1. ÁP DỤNG CHO TOÀN BỘ TRÌNH DUYỆT - ĐẢM BẢO LUÔN HIỆN DIỆN */
    /* Độ rộng cực lớn (24px) để dễ bấm trên PC */
    ::-webkit-scrollbar {
        width: 24px !important;
        height: 24px !important;
        display: block !important;
    }

    /* 2. PHẦN RÃNH TRƯỢT (TRACK) */
    ::-webkit-scrollbar-track {
        background: #1e1e1e !important; /* Nền tối để làm nổi bật thanh trượt */
        border-left: 2px solid #333 !important;
    }

    /* 3. THANH TRƯỢT CHÍNH (THUMB) - LUÔN LÀ MÀU VÀNG */
    ::-webkit-scrollbar-thumb {
        background-color: #FFD700 !important; /* Vàng Gold nguyên bản */
        background-image: linear-gradient(45deg, #FFD700 25%, #ffeb3b 50%, #FFD700 75%) !important;
        border-radius: 4px !important;
        border: 4px solid #1e1e1e !important; /* Tạo khoảng trống để thanh vàng nổi hẳn lên */
        box-shadow: 0 0 10px rgba(255, 215, 0, 0.8) !important; /* Hiệu ứng phát sáng nhẹ để dễ thấy */
    }

    /* 4. GIỮ MÀU KHI DI CHUỘT HOẶC KÉO (HOVER/ACTIVE) */
    ::-webkit-scrollbar-thumb:hover, 
    ::-webkit-scrollbar-thumb:active {
        background-color: #ffffff !important; /* Đổi sang Trắng khi đang chọn để phản hồi trực quan */
        background-image: none !important;
        box-shadow: 0 0 20px rgba(255, 255, 255, 1) !important;
    }

    /* 5. FIX ĐẶC BIỆT CHO STREAMLIT (Vùng Sidebar và Main Content) */
    [data-testid="stSidebar"]::-webkit-scrollbar,
    .main::-webkit-scrollbar,
    .stApp::-webkit-scrollbar {
        width: 24px !important;
    }
    
    /* Đối với Firefox */
    * {
        scrollbar-width: thick !important;
        scrollbar-color: #FFD700 #1e1e1e !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# CSS: Canh giữa 2 button kết quả, chữ không xuống dòng trên PC lẫn mobile
st.markdown(
    """
    <style>
    /* Canh giữa cả 2 column chứa button kết quả */
    div[data-testid="stHorizontalBlock"] {
        justify-content: center !important;
    }

    /* Mỗi column co lại vừa đủ với nội dung button */
    div[data-testid="stHorizontalBlock"] > div[data-testid="stColumn"] {
        flex: 0 1 auto !important;
        min-width: 0 !important;
        width: auto !important;
    }

    /* Button: chữ không xuống dòng, padding đẹp */
    div[data-testid="stHorizontalBlock"] div[data-testid="stColumn"] button[kind="secondary"],
    div[data-testid="stHorizontalBlock"] div[data-testid="stColumn"] button[kind="primary"] {
        white-space: nowrap !important;
        width: auto !important;
        min-width: 160px;
        padding: 0.55rem 1.4rem !important;
        font-size: clamp(13px, 2.5vw, 16px) !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)



# ⚙️ HÀM HỖ TRỢ VÀ FILE I/O
# ====================================================
def clean_text(s: str) -> str:
    if s is None:
        return ""
    
    # GIỮ NGUYÊN các pattern điền chỗ trống:
    # - 2-10 dấu chấm (có thể có space xen kẽ): .... hoặc . . . .
    # - 2-10 gạch dưới (có thể có space xen kẽ): ____ hoặc __ __
    # - Ngoặc chứa các ký tự trên: (____) hoặc (__  __) → chuẩn hóa thành (____) 
    
    temp_s = s
    placeholders = {}
    counter = 0
    
    # BƯỚC 1: Xử lý ngoặc có nhiều space/ký tự → chuẩn hóa thành 4 spaces
    # VD: (__           __) → (____)
    temp_s = re.sub(r'\([\s._-]{2,}\)', '(    )', temp_s)  # Ngoặc đơn
    temp_s = re.sub(r'\[[\s._-]{2,}\]', '[    ]', temp_s)  # Ngoặc vuông
    
    # BƯỚC 2: Lưu các pattern điền chỗ trống còn lại
    standalone_patterns = [
        r'(?<!\S)([._])(?:\s*\1){1,9}(?!\S)',  # 2-10 dấu . hoặc _ liên tiếp (có thể có space)
        r'-{2,10}',  # 2-10 gạch ngang liên tiếp
        r'\([\s]{2,}\)',  # Ngoặc đơn có spaces (đã chuẩn hóa ở bước 1)
        r'\[[\s]{2,}\]',  # Ngoặc vuông có spaces
    ]
    
    for pattern in standalone_patterns:
        # Sử dụng finditer và thay thế từng match một cách an toàn hơn
        matches = list(re.finditer(pattern, temp_s))
        # Duyệt ngược để không làm thay đổi index của các match chưa xử lý
        for match in reversed(matches):
            matched_text = match.group()
            # Sử dụng ký tự đặc biệt không trùng lặp để làm placeholder
            placeholder = f"@@@PH_{counter}@@@"
            placeholders[placeholder] = matched_text
            # Thay thế chính xác tại vị trí match
            temp_s = temp_s[:match.start()] + placeholder + temp_s[match.end():]
            counter += 1
    
    # BƯỚC 2.5: Xóa ký hiệu đáp án (*) và khoảng trắng thừa để so sánh
    temp_s = temp_s.replace("(*)", "").strip()
    temp_s = re.sub(r'\s+', ' ', temp_s).strip()
    
    # BƯỚC 3: Xóa khoảng trắng thừa (2+ spaces → 1 space)
    # temp_s = re.sub(r'\s{2,}', ' ', temp_s) # Đã gộp ở trên
    
    # BƯỚC 4: Khôi phục các pattern đã lưu
    for placeholder, original in placeholders.items():
        temp_s = temp_s.replace(placeholder, original)
    
    return temp_s.strip()

def strip_question_number(text: str) -> str:
    """Xóa số thứ tự ở đầu câu hỏi (ví dụ: '1. ', '1.1. ', '2) ') để tránh lặp số."""
    # Regex hỗ trợ cả 1., 1.1, 1.1., 1) ... theo sau là ít nhất 1 khoảng trắng
    return re.sub(r'^\s*\d+([\.\d]+)*[\.\)]?\s+', '', text).strip()

def find_file_path(source):
    """Hàm tìm đường dẫn file với cơ chế tìm kiếm đa dạng."""
    paths = [
        os.path.join(os.path.dirname(__file__), source),
        source,
        f"pages/{source}"
    ]
    for path in paths:
        if os.path.exists(path) and os.path.getsize(path) > 0:
            return path
    return None

def read_docx_paragraphs(source):
    """
    Hàm đọc paragraphs chỉ lấy TEXT (sử dụng cho cabbank, lawbank, PL1)
    """
    path = find_file_path(source)
    if not path:
        print(f"Lỗi không tìm thấy file DOCX: {source}")
        return []
    
    try:
        doc = Document(path)
        return [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    except Exception as e:
        print(f"Lỗi đọc file DOCX (chỉ text): {source}. Chi tiết: {e}")
        return []

# HÀM ĐỌC FILE MỚI: DÙNG CHO PL2 (CHỈ LẤY TEXT)
def read_pl2_data(source):
    """
    Hàm đọc paragraphs chỉ lấy TEXT (tương tự read_docx_paragraphs),
    để parse_pl2 có thể dùng logic (*).
    """
    path = find_file_path(source)
    if not path:
        print(f"Lỗi không tìm thấy file DOCX: {source}")
        return []
    
    data = []
    
    try:
        doc = Document(path)
    except Exception as e:
        print(f"Lỗi đọc file DOCX (chỉ text): {source}. Chi tiết: {e}")
        return []

    for p in doc.paragraphs:
        p_text_stripped = p.text.strip()
        if not p_text_stripped:
            continue
        
        # BỎ LOGIC HIGHLIGHT VÀNG, CHỈ LẤY TEXT VÀ ĐẶT CỜ HIGHLIGHT = FALSE
        data.append({
            "full_text": p_text_stripped,
            "has_yellow_highlight": False 
        })
        
    return data

def get_base64_encoded_file(file_path):
    fallback_base64 = "iVBORw0KGgoAAAANSUhEUAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNkYAAAAAYAAjCB0C8AAAAASUVORK5CYII="
    path_to_check = find_file_path(file_path)
    if not path_to_check:
        return fallback_base64
        
    try:
        with open(path_to_check, "rb") as f:
            return base64.b64encode(f.read()).decode("utf-8")
    except Exception as e:
        print(f"Lỗi đọc file ảnh {file_path}: {e}")
        return fallback_base64

# ====================================================
# 🌐 HÀM DỊCH THUẬT (ĐÃ CẬP NHẬT DÙNG deep_translator)
# ====================================================

# Thay thế import
from deep_translator import GoogleTranslator

@st.cache_resource
def get_translator():
    """Khởi tạo Translator với deep_translator"""
    try:
        return GoogleTranslator(source='auto', target='vi')
    except Exception as e:
        print(f"Lỗi khởi tạo translator: {e}")
        return None

# HÀM MỚI: Dịch văn bản thuần túy (Dùng cho đoạn văn)
def translate_passage_content(text):
    """
    Dịch văn bản thuần túy và cố gắng bảo toàn định dạng xuống dòng.
    """
    translator = get_translator()
    if translator is None or not text.strip():
        return f"**[LỖI]** Không thể khởi tạo translator." if not text.strip() else ""
    try:
        # Dịch nguyên khối, deep_translator thường bảo toàn line breaks nếu input có
        translated_text = translator.translate(text)
        return translated_text
    except Exception as e:
        print(f"Lỗi dịch thuật passage: {e}")
        return f"**[LỖI DỊCH THUẬT ĐOẠN VĂN]** Không thể dịch nội dung. Chi tiết: {type(e).__name__}"


# HÀM MỚI: Dùng để xây dựng chuỗi dịch cho Q/A
def build_translation_text_for_qa(q):
    """Xây dựng chuỗi văn bản đầy đủ để gửi đi dịch (chỉ Question và Options)."""
    question_text = q['question']
    options_text = '; '.join(q['options'])
    return f"Câu hỏi: {question_text}\nĐáp án: {options_text}"


# HÀM GỐC: Đã được đổi tên thành `translate_question_and_options`
def translate_question_and_options(text):
    """
    Dịch câu hỏi và đáp án sử dụng deep_translator.
    (Input là chuỗi đã được build_translation_text_for_qa định dạng)
    """
    translator = get_translator()
    
    if translator is None:
        return f"**[LỖI]** Không thể khởi tạo translator."
    
    try:
        # 1. Logic dịch Options (Dùng chung)
        def _translate_options(options_raw_text):
            a_translated_list = []
            options = [opt.strip() for opt in options_raw_text.split(';') if opt.strip()]
            for i, option_content in enumerate(options):
                if not option_content: a_translated_list.append(""); continue
                
                original_prefix_match = re.match(r'^([a-d]\.|\s*)\s*', option_content, re.IGNORECASE)
                original_prefix_with_space = original_prefix_match.group(0) if original_prefix_match else ""
                original_prefix = original_prefix_with_space.strip() if original_prefix_with_space.strip() else f"{i+1}."
                
                content_to_translate = option_content[len(original_prefix_with_space):].strip()
                if not content_to_translate: a_translated_list.append(original_prefix); continue
                
                translated_text = translator.translate(content_to_translate)
                stripped_translated_text = translated_text.strip()
                
                if stripped_translated_text.lower().startswith("một "): stripped_translated_text = stripped_translated_text[len("một "):]
                stripped_translated_text = re.sub(r'^\s*([a-d]\.|\d+\.)\s*', '', stripped_translated_text, flags=re.IGNORECASE).strip()
                if not stripped_translated_text: stripped_translated_text = translated_text.strip()
                
                a_translated_list.append(f"{original_prefix} {stripped_translated_text}")
            
            return "\n".join([f"- {opt}" for opt in a_translated_list])
        # --------------------------------------------------

        # Tách Câu hỏi và Đáp án từ input text
        q_parts = text.split('\nĐáp án: ')
        q_content = q_parts[0].replace('Câu hỏi: ', '').strip()
        a_content_raw = q_parts[1].strip() if len(q_parts) > 1 else ""
        
        q_translated = translator.translate(q_content)
        a_translated_text = _translate_options(a_content_raw)
        
        return f"**[Bản dịch Tiếng Việt]**\n\n- **Câu hỏi:** {q_translated}\n- **Các đáp án:** \n{a_translated_text}"
        
    except Exception as e:
        print(f"Lỗi dịch thuật: {e}")
        return f"**[LỖI DỊCH THUẬT]**\n- Không thể dịch nội dung. Chi tiết: {type(e).__name__}\n- Câu hỏi gốc:\n{text}"

# Đặt lại tên hàm cũ (translate_text) để tương thích với các hàm hiển thị
translate_text = translate_question_and_options
# ====================================================

# ====================================================
# 🧩 PARSER 1: NGÂN HÀNG KỸ THUẬT (CABBANK)
# ====================================================
# ... (parse_cabbank remains unchanged)
def parse_cabbank(source):
    """
    Parser cho định dạng CABBANK (Dùng dấu * trước option đúng)
    """
    paras = read_docx_paragraphs(source)
    if not paras: return []

    questions = []
    current = {"question": "", "options": [], "answer": ""}
    opt_pat = re.compile(r'(?P<star>\*)?\s*(?P<letter>[A-Da-d])[\.\)]\s+')

    for p in paras:
        matches = list(opt_pat.finditer(p))
        if not matches:
            if current["options"]:
                if current["question"] and current["options"]:
                    if not current["answer"] and current["options"]:
                        current["answer"] = current["options"][0]
                    questions.append(current)
                # Xóa số thứ tự cho câu hỏi mới
                current = {"question": strip_question_number(clean_text(p)), "options": [], "answer": ""}
            else:
                if not current["question"]:
                    current["question"] = strip_question_number(clean_text(p))
                else:
                    current["question"] += " " + clean_text(p)
            continue

        pre_text = p[:matches[0].start()].strip()
        if pre_text:
            if current["options"]:
                if current["question"] and current["options"]:
                    if not current["answer"] and current["options"]:
                        current["answer"] = current["options"][0]
                    questions.append(current)
                current = {"question": strip_question_number(clean_text(pre_text)), "options": [], "answer": ""}
            else:
                if not current["question"]:
                    current["question"] = strip_question_number(clean_text(pre_text))
                else:
                    current["question"] += " " + clean_text(pre_text)

        for i, m in enumerate(matches):
            s = m.end()
            e = matches[i + 1].start() if i + 1 < len(matches) else len(p)
            raw_body = p[s:e]
            # Nhận diện (*) ở cuối đáp án (ví dụ: "A. text (*)")
            has_star_end = bool(re.search(r'\(\*\)\s*$', raw_body.strip()))
            opt_body = clean_text(raw_body)
            letter = m.group('letter').lower()
            opt = f"{letter}. {opt_body}"
            current["options"].append(opt)
            if m.group("star") or has_star_end: current["answer"] = opt

    if current["question"] and current["options"]:
        if not current["answer"] and current["options"]:
            current["answer"] = current["options"][0]
        questions.append(current)
    return questions
# ====================================================
# 🧩 PARSER 2: NGÂN HÀNG LUẬT (LAWBANK)
# ====================================================
# ... (parse_lawbank remains unchanged)
def parse_lawbank(source):
    """
    Parser cho định dạng LAWBANK (Dùng dấu * trước option đúng)
    """
    paras = read_docx_paragraphs(source)
    if not paras: return []

    questions = []
    current = {"question": "", "options": [], "answer": ""}
    opt_pat = re.compile(r'(?<![A-Za-z0-9/])(?P<star>\*)?\s*(?P<letter>[A-Da-d])[\.\)]\s+')

    for p in paras:
        if re.match(r'^\s*Ref', p, re.I): continue
        matches = list(opt_pat.finditer(p))
        
        if not matches:
            if current["options"]:
                if current["question"] and current["options"]:
                    if not current["answer"] and current["options"]:
                        current["answer"] = current["options"][0]
                    questions.append(current)
                # Xóa số thứ tự cho câu hỏi mới
                current = {"question": strip_question_number(clean_text(p)), "options": [], "answer": ""}
            else:
                if not current["question"]:
                    current["question"] = strip_question_number(clean_text(p))
                else:
                    current["question"] += " " + clean_text(p)
            continue

        first_match = matches[0]
        pre_text = p[:first_match.start()].strip()
        if pre_text:
            if current["options"]:
                if current["question"] and current["options"]:
                    if not current["answer"] and current["options"]:
                        current["answer"] = current["options"][0]
                    questions.append(current)
                current = {"question": strip_question_number(clean_text(pre_text)), "options": [], "answer": ""}
            else:
                if not current["question"]:
                    current["question"] = strip_question_number(clean_text(pre_text))
                else:
                    current["question"] += " " + clean_text(pre_text)

        for i, m in enumerate(matches):
            s = m.end()
            e = matches[i+1].start() if i+1 < len(matches) else len(p)
            raw_body = p[s:e]
            # Nhận diện (*) ở cuối đáp án (ví dụ: "A. text (*)")
            has_star_end = bool(re.search(r'\(\*\)\s*$', raw_body.strip()))
            opt_body = clean_text(raw_body)
            letter = m.group("letter").lower()
            option = f"{letter}. {opt_body}"
            current["options"].append(option)
            if m.group("star") or has_star_end: current["answer"] = option

    if current["question"] and current["options"]:
        if not current["answer"] and current["options"]:
            current["answer"] = current["options"][0]
        questions.append(current)
    return questions

# ====================================================
# 🧩 PARSER 3: PHỤ LỤC 1 (Dùng dấu (*))
# ====================================================
# ... (parse_pl1 remains unchanged)
def parse_pl1(source):
    """
    Parser cho định dạng PL1 (sử dụng dấu (*) để nhận diện đáp án đúng)
    """
    paras = read_docx_paragraphs(source)
    if not paras: return []

    questions = []
    current = {"question": "", "options": [], "answer": ""}
    
    q_start_pat = re.compile(r'^\s*(\d+)[\.\)]\s*') 
    phrase_start_pat = re.compile(r'Choose the correct group of words', re.I)
    opt_prefix_pat = re.compile(r'^\s*[A-Ca-c]([\.\)]|\s+)\s*') 
    labels = ["a", "b", "c"]
    MAX_OPTIONS = 3

    def finalize_current_question(q_dict, q_list):
        if q_dict["question"] and q_dict["options"]:
            q_list.append(q_dict)
        return {"question": "", "options": [], "answer": ""}
    
    for p in paras:
        clean_p = p.strip()
        if not clean_p: continue
        
        is_q_start_phrased = phrase_start_pat.search(clean_p)
        is_explicitly_numbered = q_start_pat.match(clean_p) 
        is_max_options_reached = len(current["options"]) >= MAX_OPTIONS
        is_question_started = current["question"]
        is_first_line = not is_question_started and not current["options"]
        
        must_switch_q = (
            is_first_line or                            
            is_q_start_phrased or                       
            (is_question_started and is_max_options_reached)
        )
        
        if must_switch_q:
            current = finalize_current_question(current, questions)
            q_text = strip_question_number(clean_p)
            current["question"] = q_text
            
        else:
            if is_question_started and not is_max_options_reached:
                is_correct = False
                
                # SỬ DỤNG DẤU (*)
                if "(*)" in clean_p:
                    is_correct = True
                    clean_p = clean_p.replace("(*)", "").strip() 
                
                match_prefix = opt_prefix_pat.match(clean_p)
                if match_prefix:
                    clean_p = clean_p[match_prefix.end():].strip()
                    
                idx = len(current["options"])
                    
                if idx < len(labels):
                    label = labels[idx]
                    opt_text = f"{label}. {clean_p}"
                    current["options"].append(opt_text)
                    
                    if is_correct:
                        current["answer"] = opt_text
            
            elif is_question_started:
                 current["question"] += " " + clean_p
        
            elif not is_question_started and not current["options"]:
                current["question"] = clean_p

    current = finalize_current_question(current, questions)
        
    return questions

# ====================================================
# 🧩 PARSER 4: PHỤ LỤC 2 (Dùng dấu (*))
# ====================================================
# ... (parse_pl2 remains unchanged)
def parse_pl2(source):
    """
    Parser cho định dạng PL2 (Sử dụng ký hiệu (*) để nhận diện đáp án đúng)
    """
    data = read_pl2_data(source) # SỬ DỤNG HÀM ĐỌC ĐÃ SỬA CHỈ LẤY TEXT
    if not data: return []

    questions = []
    current = {"question": "", "options": [], "answer": ""}
    
    q_start_pat = re.compile(r'^\s*(\d+)[\.\)]\s*') 
    phrase_start_pat = re.compile(r'Choose the correct group of words', re.I)
    opt_prefix_pat = re.compile(r'^\s*[A-Ca-c]([\.\)]|\s+)\s*') 
    labels = ["a", "b", "c"]
    MAX_OPTIONS = 3

    def finalize_current_question(q_dict, q_list):
        if q_dict["question"] and q_dict["options"]:
            q_list.append(q_dict)
        return {"question": "", "options": [], "answer": ""}
    
    for p_data in data:
        clean_p = p_data["full_text"].strip()
        if not clean_p: continue
        
        is_q_start_phrased = phrase_start_pat.search(clean_p)
        is_explicitly_numbered = q_start_pat.match(clean_p) 
        is_max_options_reached = len(current["options"]) >= MAX_OPTIONS
        is_question_started = current["question"]
        is_first_line = not is_question_started and not current["options"]
        
        must_switch_q = (
            is_first_line or                            
            is_q_start_phrased or                       
            (is_question_started and is_max_options_reached)
        )
        
        if must_switch_q:
            current = finalize_current_question(current, questions)
            q_text = strip_question_number(clean_p)
            current["question"] = q_text
            
        else:
            if is_question_started and not is_max_options_reached:
                is_correct = False
                
                # SỬ DỤNG LOGIC DẤU (*)
                if "(*)" in clean_p:
                    is_correct = True
                    clean_p = clean_p.replace("(*)", "").strip() # Loại bỏ ký hiệu sau khi phát hiện
                
                match_prefix = opt_prefix_pat.match(clean_p)
                if match_prefix:
                    clean_p = clean_p[match_prefix.end():].strip()
                    
                idx = len(current["options"])
                if idx < len(labels):
                    label = labels[idx]
                    opt_text = f"{label}. {clean_p}"
                    current["options"].append(opt_text)
                    
                    if is_correct:
                        current["answer"] = opt_text
            
            elif is_question_started:
                 current["question"] += " " + clean_p
        
            elif not is_question_started and not current["options"]:
                current["question"] = clean_p

    current = finalize_current_question(current, questions)
        
    return questions

# ====================================================
# 🧩 PARSER 5: PHỤ LỤC 3 - BÀI ĐỌC HIỂU (PASSAGE-BASED) - ĐÃ SỬA LỖI PARAGRAPH 2
# ====================================================
# ... (parse_pl3_passage_bank remains unchanged)
def parse_pl3_passage_bank(source):
    """
    Parser cho định dạng PL3 (Bài đọc hiểu)
    - Fix: Xử lý đúng cho câu hỏi điền chỗ trống (Paragraph 2) bằng cách tạo câu hỏi tường minh.
    - Fix: Xử lý các đoạn văn không có câu hỏi.
    - Fix: Giảm khả năng nhận diện sai câu hỏi (yêu cầu có khoảng trắng sau số thứ tự).
    """
    path = find_file_path(source)
    if not path:
        print(f"Lỗi không tìm thấy file DOCX: {source}")
        return []
    
    questions = []
    current_group = None
    group_content = ""
    current_q_num = 0
    
    paragraph_start_pat = re.compile(r'^\s*Paragraph\s*(\d+)\s*\.\s*', re.I)
    q_start_pat = re.compile(r'^\s*(?P<q_num>\d+)\s*[\.\)]\s+', re.I)
    opt_pat_single = re.compile(r'^\s*(?P<letter>[A-Da-d])[\.\)]\s*(?P<text>.*?)(\s*\(\*\))?$', re.I)
    
    try:
        doc = Document(path)
    except Exception as e:
        print(f"Lỗi đọc file DOCX: {source}. Chi tiết: {e}")
        return []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text: continue
        
        is_new_paragraph_group = paragraph_start_pat.match(text)
        match_q_start = q_start_pat.match(text)
        
        if is_new_paragraph_group:
            if current_group is not None:
                if current_group.get('question'):
                    questions.append(current_group)
                elif group_content.strip():
                    questions.append({
                        'group_name': current_group['group_name'],
                        'paragraph_content': group_content.strip(),
                        'question': '', 'options': {}, 'correct_answer': '', 'number': 0,
                        'is_content_only': True
                    })
            
            group_name = is_new_paragraph_group.group(0).strip()
            current_group = {
                'group_name': group_name,
                'paragraph_content': "", 'question': "", 'options': {}, 'correct_answer': "", 'number': 0
            }
            group_content = paragraph.text + "\n"
            current_q_num = 0
            continue
            
        if current_group is None:
            continue
            
        if match_q_start:
            if current_group.get('question') and current_group.get('options'):
                 questions.append(current_group)
            
            q_num_str = match_q_start.group('q_num')
            remaining_text = text[match_q_start.end():].strip()
            
            is_fill_in_blank = bool(re.search(r'\(\s*\d+\s*\)', group_content))
            
            if is_fill_in_blank:
                q_text = f"Chọn đáp án thích hợp cho ô trống **({q_num_str})** trong đoạn văn trên."
                first_option_text = remaining_text
            else:
                q_text = remaining_text
                first_option_text = ""
            
            current_group = {
                'group_name': current_group['group_name'],
                'paragraph_content': group_content.strip(), 
                'question': clean_text(q_text),
                'options': {}, 'correct_answer': "",
                'number': int(q_num_str) 
            }
            current_q_num = int(q_num_str)
            
            if is_fill_in_blank and first_option_text:
                match_opt = opt_pat_single.match(first_option_text)
                if match_opt:
                    letter = match_opt.group('letter').upper()
                    opt_text_raw = match_opt.group('text').strip()
                    is_correct = match_opt.group(3) is not None
                    opt_text = clean_text(opt_text_raw.replace("(*)", "").strip())
                    full_opt_text = f"{letter}. {opt_text}"
                    current_group['options'][letter] = full_opt_text
                    if is_correct:
                        current_group['correct_answer'] = letter
            
        elif current_q_num > 0:
            match_opt = opt_pat_single.match(text)
            if match_opt:
                letter = match_opt.group('letter').upper()
                opt_text_raw = match_opt.group('text').strip()
                is_correct = match_opt.group(3) is not None
                opt_text = clean_text(opt_text_raw.replace("(*)", "").strip())
                full_opt_text = f"{letter}. {opt_text}"
                current_group['options'][letter] = full_opt_text
                if is_correct:
                    current_group['correct_answer'] = letter
            else:
                current_group['question'] += " " + clean_text(text)
                
        elif current_group is not None and not is_new_paragraph_group:
            if current_q_num == 0:
                group_content += paragraph.text + "\n"
        
    if current_group is not None:
        if current_group.get('question'):
            questions.append(current_group)
        elif group_content.strip():
            questions.append({
                'group_name': current_group['group_name'],
                'paragraph_content': group_content.strip(),
                'question': '', 'options': {}, 'correct_answer': '', 'number': 0,
                'is_content_only': True
            })

    final_questions = []
    global_q_counter = 1 
    for q in questions:
        if q.get('is_content_only'):
            final_questions.append({
                'question': '', 'options': [], 'answer': '', 'number': 0, 'global_number': 0,
                'group': q['group_name'], 
                'paragraph_content': q['paragraph_content'],
                'is_content_only': True
            })
            continue

        if not q.get('correct_answer') and len(q.get('options', {})) > 0:
             # Theo yêu cầu mới, không tự động gán đáp án đầu tiên nếu không tìm thấy (*)
             pass
        
        if not q.get('correct_answer') or not q.get('options'):
            # Vẫn cho phép câu hỏi hiển thị nhưng không có đáp án đúng nếu muốn, 
            # hoặc bỏ qua nếu dữ liệu lỗi. Ở đây ta bỏ qua để đảm bảo chất lượng.
            continue
        
        options_list = list(q['options'].values()) 
        
        final_questions.append({
            'question': q['question'],
            'options': options_list, 
            'answer': q['options'][q['correct_answer']],
            'number': q['number'],
            'global_number': global_q_counter,
            'group': q['group_name'], 
            'paragraph_content': q['paragraph_content']
        })
        global_q_counter += 1

    return final_questions
def parse_pl4_law_process(source):
    path = find_file_path(source)
    if not path: return []
    
    questions = []
    current_group = None
    group_content = ""
    local_q_counter = 0 
    group_name = "Paragraph 1"
    
    # Pattern nhận diện: Paragraph X, Số câu hỏi, và Đáp án (*)
    paragraph_start_pat = re.compile(r'^\s*Paragraph\s*(\d+)\s*\.\s*', re.I)
    q_start_pat = re.compile(r'^\s*(?P<q_num>\d+)\s*[\.\)]\s*', re.I)
    opt_pat_single = re.compile(r'^\s*(?P<letter>[A-Da-d])[\.\)]\s*(?P<text>.*?)(\s*\(\*\))?$', re.I)
    # Pattern nhận diện vị trí điền trống trong đoạn văn: (1), (2)...
    blank_pattern = re.compile(r'\(\s*(\d+)\s*\)')

    try:
        doc = Document(path)
    except Exception as e:
        print(f"Lỗi: {e}")
        return []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text: continue
        
        is_new_paragraph_group = paragraph_start_pat.match(text)
        match_q_start = q_start_pat.match(text)
        
        if is_new_paragraph_group:
            if current_group and current_group.get('question'):
                questions.append(current_group)
            group_name = is_new_paragraph_group.group(0).strip()
            current_group = None
            group_content = ""
            continue

        if match_q_start:
            if current_group and current_group.get('question'):
                questions.append(current_group)
            
            q_num_str = match_q_start.group('q_num')
            remaining_text = text[match_q_start.end():].strip()
            
            # KIỂM TRA: Nếu trong đoạn văn có ký hiệu (1), (2)... thì đây là câu hỏi điền trống
            is_fill_blank = bool(blank_pattern.search(group_content))
            
            if is_fill_blank:
                # Tạo câu hỏi miêu tả tự động
                q_text = f"Chọn đáp án thích hợp cho vị trí trống ………… ({q_num_str}) trong đoạn văn trên."
                first_opt_raw = remaining_text # Text còn lại sau số thứ tự chính là Option A
            else:
                q_text = remaining_text
                first_opt_raw = ""

            current_group = {
                'group_name': group_name,
                'paragraph_content': group_content.strip(),
                'question': clean_text(q_text),
                'options': {},
                'correct_answer': "",
                'number': int(q_num_str)
            }

            # Nếu là điền trống, xử lý Option A ngay tại dòng này
            if is_fill_blank and first_opt_raw:
                m_opt = opt_pat_single.match(first_opt_raw)
                if m_opt:
                    letter = m_opt.group('letter').upper()
                    is_correct = m_opt.group(3) is not None
                    opt_body = clean_text(m_opt.group('text').replace("(*)", "").strip())
                    current_group['options'][letter] = f"{letter}. {opt_body}"
                    if is_correct: current_group['correct_answer'] = letter
            continue

        if current_group:
            match_opt = opt_pat_single.match(text)
            if match_opt:
                letter = match_opt.group('letter').upper()
                is_correct = match_opt.group(3) is not None
                opt_body = clean_text(match_opt.group('text').replace("(*)", "").strip())
                current_group['options'][letter] = f"{letter}. {opt_body}"
                if is_correct: current_group['correct_answer'] = letter
            else:
                # Nếu không phải option, nối tiếp vào câu hỏi (trường hợp câu hỏi dài nhiều dòng)
                current_group['question'] += " " + clean_text(text)
        else:
            # Thu thập nội dung đoạn văn
            group_content += text + "\n"

    if current_group and current_group.get('question'):
        questions.append(current_group)

    # Chuyển đổi sang định dạng chuẩn để hiển thị
    final_questions = []
    for q in questions:
        if not q.get('options'): continue
        
        # Nếu không có đáp án đúng, mặc định chọn đáp án đầu tiên (hoặc xử lý khác tùy bạn)
        # Ở đây ta nên kiểm tra nếu correct_answer rỗng thì có thể log lỗi hoặc bỏ qua
        if not q.get('correct_answer'):
            # Vẫn để fallback này nếu cần, nhưng chỉ khi thật sự không tìm thấy (*)
            # Tuy nhiên, theo yêu cầu người dùng, ta nên cẩn thận với việc tự set câu a
            pass
            
        final_questions.append({
            'question': q['question'],
            'options': list(q['options'].values()),
            'answer': q['options'].get(q['correct_answer'], list(q['options'].values())[0]) if q.get('correct_answer') else list(q['options'].values())[0],
            'number': q['number'], 
            'group': q['group_name'],
            'paragraph_content': q['paragraph_content']
        })
    return final_questions
def parse_pl5_specialized(source):
    """
    Parser cho định dạng PL5 (Chuyên ngành)
    - Câu hỏi bắt đầu bằng số: 1., 2., 3., ...
    - Đáp án A., B., C. với câu đúng có dấu (*) ở cuối
    """
    paras = read_docx_paragraphs(source)
    if not paras: return []

    questions = []
    current = {"question": "", "options": [], "answer": ""}
    
    # Pattern nhận diện số thứ tự câu hỏi
    q_start_pat = re.compile(r'^\s*(\d+)\s*[\.\)]\s*')
    # Pattern nhận diện đáp án A, B, C, D, E (với hoặc không có dấu (*) ở cuối)
    # Hỗ trợ nhiều loại dấu ngăn cách: . ) / -
    opt_pat = re.compile(r'^\s*(?P<letter>[A-Ea-e])[\.\)\/\-]\s*(?P<text>.*?)\s*(?P<star>\(\*\))?\s*$', re.I)
    
    for p in paras:
        clean_p = p.strip()
        if not clean_p: continue
        
        # Kiểm tra xem có phải câu hỏi mới không (Ví dụ: "1.", "Câu 1:", "1.1.")
        q_match = q_start_pat.match(clean_p)
        if q_match:
            # Lưu câu hỏi cũ nếu có
            if current["question"] and current["options"]:
                questions.append(current)
            
            # Bắt đầu câu hỏi mới
            q_text = strip_question_number(clean_p)
            current = {"question": q_text, "options": [], "answer": ""}
            continue
        
        # Kiểm tra xem có phải đáp án không
        opt_match = opt_pat.match(clean_p)
        if opt_match and current["question"]:
            letter = opt_match.group("letter").lower()
            opt_text = opt_match.group("text").strip()
            has_star = opt_match.group("star") is not None
            
            # Đảm bảo opt_text không còn dấu (*) nếu regex chưa bắt hết
            opt_text = opt_text.replace("(*)", "").strip()
            
            full_option = f"{letter}. {opt_text}"
            current["options"].append(full_option)
            
            if has_star:
                current["answer"] = full_option
        else:
            # Nếu không phải câu hỏi mới hoặc đáp án, nối vào câu hỏi hiện tại
            # Nhưng bỏ qua nếu là các dòng rác như "Trang X", "Page X", "---"
            junk_patterns = [r'^Trang\s+\d+', r'^Page\s+\d+', r'^-+$', r'^_+$']
            if any(re.match(jp, clean_p, re.I) for jp in junk_patterns):
                continue

            if current["question"]:
                # Tránh lặp lại nội dung nếu đã có
                if clean_text(clean_p) not in current["question"]:
                    current["question"] += " " + clean_text(clean_p)
    
    # Lưu câu hỏi cuối cùng
    if current["question"] and current["options"]:
        questions.append(current)
    
    return questions
# ====================================================
# 🌟 HÀM: LOGIC DỊCH ĐỘC QUYỀN (EXCLUSIVE TRANSLATION)
# ====================================================
if 'active_translation_key' not in st.session_state: st.session_state.active_translation_key = None
# Thêm trạng thái cho dịch đoạn văn
if 'active_passage_translation' not in st.session_state: st.session_state.active_passage_translation = None
if 'passage_translations_cache' not in st.session_state: st.session_state.passage_translations_cache = {}

def on_translate_toggle(key_clicked):
    """Callback function để quản lý chế độ Dịch ĐỘC QUYỀN (Q&A)."""
    toggle_key = f"toggle_{key_clicked}"
    # Check the state of the toggle in session state (it is the state *after* the click)
    is_on_after_click = st.session_state.get(toggle_key, False)
    
    if is_on_after_click:
        # User turned this specific toggle ON -> Make it the active key
        st.session_state.active_translation_key = key_clicked
    elif st.session_state.active_translation_key == key_clicked:
        # User turned this specific toggle OFF -> Clear the active key
        st.session_state.active_translation_key = None
    
def on_passage_translate_toggle(passage_id_clicked):
    """Callback function để quản lý chế độ Dịch ĐỘC QUYỀN (Passage)."""
    toggle_key = f"toggle_passage_{passage_id_clicked}"
    is_on_after_click = st.session_state.get(toggle_key, False)

    if is_on_after_click:
        # User turned this specific toggle ON -> Make it the active passage key
        st.session_state.active_passage_translation = passage_id_clicked
    elif st.session_state.active_passage_translation == passage_id_clicked:
        # User turned this specific toggle OFF -> Clear the active key
        st.session_state.active_passage_translation = None

# ====================================================
# 🌟 HÀM: XEM TOÀN BỘ CÂU HỎI (CẬP NHẬT CHỨC NĂNG DỊCH)
# ====================================================
def display_all_questions(questions):
    st.markdown('<div class="result-title"><h3>📚 TOÀN BỘ NGÂN HÀNG CÂU HỎI</h3></div>', unsafe_allow_html=True)
    if not questions:
        st.warning("Không có câu hỏi nào để hiển thị.")
        return

    # Khởi tạo trạng thái submitted cho chế độ "all"
    if 'all_submitted' not in st.session_state:
        st.session_state.all_submitted = False

    current_passage_id = None

    if not st.session_state.all_submitted:
        # ===== CHƯA NỘP BÀI: hiển thị radio button để chọn =====
        for i, q in enumerate(questions, start=1):
            q_key = f"all_q_{i}_{hash(q['question'])}"
            translation_key = f"trans_{q_key}"
            is_active = (translation_key == st.session_state.active_translation_key)

            # Đoạn văn (PL3/PL4)
            passage_content = q.get('paragraph_content', '').strip()
            group_name = q.get('group', '')
            if passage_content:
                passage_id = f"passage_{group_name}_{hash(passage_content)}"
                is_passage_active = (passage_id == st.session_state.active_passage_translation)
                if passage_id != current_passage_id:
                    st.markdown(f'<div class="paragraph-title">**{group_name}**</div>', unsafe_allow_html=True)
                    st.markdown(f'<div class="paragraph-content-box">{passage_content}</div>', unsafe_allow_html=True)
                    if not q.get('is_content_only'):
                        st.toggle("🌐 Dịch đoạn văn sang Tiếng Việt", value=is_passage_active,
                                  key=f"toggle_passage_{passage_id}",
                                  on_change=on_passage_translate_toggle, args=(passage_id,))
                        if is_passage_active:
                            translated_passage = st.session_state.passage_translations_cache.get(passage_id)
                            if not isinstance(translated_passage, str):
                                translated_passage = translate_passage_content(passage_content)
                                st.session_state.passage_translations_cache[passage_id] = translated_passage
                            st.markdown(f"""
                            <div data-testid="stAlert" class="stAlert stAlert-info">
                                <div style="font-size: 18px; line-height: 1.6; color: white; padding: 10px;">
                                    <strong style="color: #FFD700;">[Bản dịch Đoạn văn]</strong>
                                    <div class="paragraph-content-box" style="white-space: pre-wrap; margin-bottom: 0px; padding: 10px; background-color: rgba(0, 0, 0, 0.5); border-left: 3px solid #00d4ff;">
                                    {translated_passage}
                                    </div>
                                </div>
                            </div>""", unsafe_allow_html=True)
                    st.markdown("---")
                    current_passage_id = passage_id

            if q.get('is_content_only'):
                continue

            # Số thứ tự
            if q.get('group', '').startswith('Paragraph') and q.get('paragraph_content'):
                display_num = q.get('number', i)
            else:
                display_num = i

            st.markdown(f'<div class="bank-question-text">{display_num}. {q["question"]}</div>', unsafe_allow_html=True)

            st.toggle("🌐 Dịch Câu hỏi & Đáp án sang Tiếng Việt", value=is_active,
                      key=f"toggle_{translation_key}",
                      on_change=on_translate_toggle, args=(translation_key,))
            if is_active:
                translated_content = st.session_state.translations.get(translation_key)
                if not isinstance(translated_content, str):
                    st.session_state.translations[translation_key] = translate_text(build_translation_text_for_qa(q))
                    translated_content = st.session_state.translations[translation_key]
                st.info(translated_content, icon="🌐")

            default_val = st.session_state.get(q_key, None)
            # Ẩn dấu (*) khỏi nhãn của radio button
            options_clean = [opt.replace("(*)", "").strip() for opt in q["options"]]
            st.radio("", options_clean,
                     index=options_clean.index(default_val.replace("(*)", "").strip()) if default_val and default_val.replace("(*)", "").strip() in options_clean else None,
                     key=q_key)
            st.markdown('<div class="question-separator"></div>', unsafe_allow_html=True)

        if st.button("✅ Nộp bài", key="submit_all"):
            st.session_state.all_submitted = True
            st.session_state.active_translation_key = None
            st.session_state.active_passage_translation = None
            st.rerun()

    else:
        # ===== SAU KHI NỘP BÀI: hiển thị kết quả =====
        score = 0
        current_passage_id = None
        for i, q in enumerate(questions, start=1):
            q_key = f"all_q_{i}_{hash(q['question'])}"
            selected_opt = st.session_state.get(q_key)
            correct = clean_text(q["answer"])
            is_correct = (selected_opt is not None) and (clean_text(selected_opt) == correct)
            translation_key = f"trans_{q_key}"
            is_active = (translation_key == st.session_state.active_translation_key)

            # Đoạn văn (PL3/PL4)
            passage_content = q.get('paragraph_content', '').strip()
            group_name = q.get('group', '')
            if passage_content:
                passage_id = f"passage_{group_name}_{hash(passage_content)}"
                is_passage_active = (passage_id == st.session_state.active_passage_translation)
                if passage_id != current_passage_id:
                    st.markdown(f'<div class="paragraph-title">**{group_name}**</div>', unsafe_allow_html=True)
                    st.markdown(f'<div class="paragraph-content-box">{passage_content}</div>', unsafe_allow_html=True)
                    if not q.get('is_content_only'):
                        st.toggle("🌐 Dịch đoạn văn sang Tiếng Việt", value=is_passage_active,
                                  key=f"toggle_passage_{passage_id}",
                                  on_change=on_passage_translate_toggle, args=(passage_id,))
                        if is_passage_active:
                            translated_passage = st.session_state.passage_translations_cache.get(passage_id)
                            if not isinstance(translated_passage, str):
                                translated_passage = translate_passage_content(passage_content)
                                st.session_state.passage_translations_cache[passage_id] = translated_passage
                            st.markdown(f"""
                            <div data-testid="stAlert" class="stAlert stAlert-info">
                                <div style="font-size: 18px; line-height: 1.6; color: white; padding: 10px;">
                                    <strong style="color: #FFD700;">[Bản dịch Đoạn văn]</strong>
                                    <div class="paragraph-content-box" style="white-space: pre-wrap; margin-bottom: 0px; padding: 10px; background-color: rgba(0, 0, 0, 0.5); border-left: 3px solid #00d4ff;">
                                    {translated_passage}
                                    </div>
                                </div>
                            </div>""", unsafe_allow_html=True)
                    st.markdown("---")
                    current_passage_id = passage_id

            if q.get('is_content_only'):
                continue

            # Số thứ tự
            if q.get('group', '').startswith('Paragraph') and q.get('paragraph_content'):
                display_num = q.get('number', i)
            else:
                display_num = i

            st.markdown(f'<div class="bank-question-text">{display_num}. {q["question"]}</div>', unsafe_allow_html=True)

            st.toggle("🌐 Dịch Câu hỏi & Đáp án sang Tiếng Việt", value=is_active,
                      key=f"toggle_{translation_key}",
                      on_change=on_translate_toggle, args=(translation_key,))
            if is_active:
                translated_content = st.session_state.translations.get(translation_key)
                if not isinstance(translated_content, str):
                    st.session_state.translations[translation_key] = translate_text(build_translation_text_for_qa(q))
                    translated_content = st.session_state.translations[translation_key]
                st.info(translated_content, icon="🌐")

            # Hiển thị đáp án với màu sắc
            for opt in q["options"]:
                opt_clean = clean_text(opt)
                # Ẩn dấu (*) khi hiển thị kết quả
                opt_display = opt.replace("(*)", "").strip()
                if opt_clean == correct:
                    ans_class = "bank-answer-text answer-correct"
                elif selected_opt is not None and clean_text(selected_opt) == opt_clean:
                    ans_class = "bank-answer-text answer-selected-wrong"
                else:
                    ans_class = "bank-answer-text answer-wrong"
                st.markdown(f'<div class="{ans_class}">{opt_display}</div>', unsafe_allow_html=True)

            if is_correct:
                score += 1
            st.markdown('<div class="question-separator"></div>', unsafe_allow_html=True)

        st.markdown(f'<div class="result-title"><h3>🎯 KẾT QUẢ: {score}/{len(questions)}</h3></div>', unsafe_allow_html=True)
        if st.button("🔄 Làm lại toàn bộ", key="reset_all"):
            for i, q in enumerate(questions, start=1):
                q_key = f"all_q_{i}_{hash(q['question'])}"
                st.session_state.pop(q_key, None)
            st.session_state.all_submitted = False
            st.session_state.active_translation_key = None
            st.session_state.active_passage_translation = None
            st.rerun()

# ====================================================
# 🌟 HÀM: TEST MODE (CẬP NHẬT CHỨC NĂNG DỊCH)
# ====================================================
def get_random_questions(questions, count=50):
    if len(questions) <= count: return questions
    return random.sample(questions, count)


def build_caav_mock_exam_questions():
    """
    Tạo bộ đề thi thử CAAV: 70 câu
    - 10 câu random từ Human Factor
    - 30 câu random từ CAAV Cabin
    - 30 câu random từ CAAV Law (Module 10.1 và 10.2)
    """
    try:
        qs_hf = parse_cabbank("caav hf.docx")
        qs_cab = parse_cabbank("caav cab.docx")
        qs_law1 = parse_lawbank("caav law1.docx")
        qs_law2 = parse_lawbank("caav law2.docx")
        
        if not qs_hf or not qs_cab or not qs_law1 or not qs_law2:
            return [], "Không tìm thấy đủ dữ liệu từ các file CAAV (HF, Cabin, Law 10.1, Law 10.2)."
            
        final_hf = random.sample(qs_hf, min(10, len(qs_hf)))
        final_cab = random.sample(qs_cab, min(30, len(qs_cab)))
        
        all_law = qs_law1 + qs_law2
        final_law = random.sample(all_law, min(30, len(all_law)))
        
        final_qs = final_hf + final_cab + final_law
        random.shuffle(final_qs)
        
        for idx, q in enumerate(final_qs, start=1):
            q['global_number'] = idx
            
        return final_qs, None
    except Exception as e:
        return [], str(e)

def display_caav_mock_test_mode():
    TOTAL_QUESTIONS = 70
    PASS_SCORE_PCT = 75 # 75/100
    test_key_prefix = "caav_mock_test"
    
    if f"{test_key_prefix}_started" not in st.session_state:
        st.session_state[f"{test_key_prefix}_started"] = False
    if f"{test_key_prefix}_submitted" not in st.session_state:
        st.session_state[f"{test_key_prefix}_submitted"] = False
    if f"{test_key_prefix}_questions" not in st.session_state:
        st.session_state[f"{test_key_prefix}_questions"] = []

    if not st.session_state[f"{test_key_prefix}_started"]:
        st.markdown('<div class="result-title"><h3>🎓 THI THỬ CAAV (70 CÂU)</h3></div>', unsafe_allow_html=True)
        st.markdown(f"""
        <div style="background:rgba(255,215,0,0.08); border-left:4px solid #FFD700; padding:16px 22px; border-radius:8px; margin-bottom:20px; font-size:18px; text-align:center; line-height:2;">
        📋 Cơ cấu đề: <b>10 HF + 30 Cabin + 30 Law</b><br>
        🎯 Điểm đạt (PASS): <b>{PASS_SCORE_PCT}/100</b>
        </div>
        """, unsafe_allow_html=True)

        if st.button("🚀 Bắt đầu Thi thử CAAV", key=f"{test_key_prefix}_start_btn"):
            with st.spinner("⏳ Đang tạo đề thi ngẫu nhiên..."):
                test_qs, error = build_caav_mock_exam_questions()
            if error:
                st.error(f"❌ {error}")
                return
            st.session_state[f"{test_key_prefix}_questions"] = test_qs
            st.session_state[f"{test_key_prefix}_started"] = True
            st.session_state[f"{test_key_prefix}_submitted"] = False
            st.rerun()
        return

    test_batch = st.session_state[f"{test_key_prefix}_questions"]
    
    if not st.session_state[f"{test_key_prefix}_submitted"]:
        st.markdown(f'<div class="result-title"><h3>⏳ ĐANG THI THỬ CAAV — {len(test_batch)} CÂU</h3></div>', unsafe_allow_html=True)
        for i, q in enumerate(test_batch, start=1):
            q_key = f"{test_key_prefix}_q_{i}_{hash(q['question'])}"
            st.markdown(f'<div class="bank-question-text">{i}. {q["question"]}</div>', unsafe_allow_html=True)
            options_clean = [opt.replace("(*)", "").strip() for opt in q["options"]]
            st.radio("", options_clean, index=None, key=q_key)
            st.markdown('<div class="question-separator"></div>', unsafe_allow_html=True)
            
        if st.button("✅ Nộp bài", key=f"{test_key_prefix}_submit_btn"):
            st.session_state[f"{test_key_prefix}_submitted"] = True
            st.rerun()
    else:
        score = 0
        for i, q in enumerate(test_batch, start=1):
            q_key = f"{test_key_prefix}_q_{i}_{hash(q['question'])}"
            selected_opt = st.session_state.get(q_key)
            correct = clean_text(q["answer"])
            is_correct = (selected_opt is not None) and (clean_text(selected_opt) == correct)
            
            st.markdown(f'<div class="bank-question-text">{i}. {q["question"]}</div>', unsafe_allow_html=True)
            for opt in q["options"]:
                opt_clean = clean_text(opt)
                opt_display = opt.replace("(*)", "").strip()
                if opt_clean == correct:
                    ans_class = "bank-answer-text answer-correct"
                elif selected_opt is not None and clean_text(selected_opt) == opt_clean:
                    ans_class = "bank-answer-text answer-selected-wrong"
                else:
                    ans_class = "bank-answer-text answer-wrong"
                st.markdown(f'<div class="{ans_class}">{opt_display}</div>', unsafe_allow_html=True)

            if is_correct: score += 1
            st.markdown('<div class="question-separator"></div>', unsafe_allow_html=True)

        final_score_scaled = (score / len(test_batch)) * 100
        st.markdown(f'<div class="result-title"><h3>🎯 KẾT QUẢ: {int(final_score_scaled)}/100</h3></div>', unsafe_allow_html=True)
        if final_score_scaled >= PASS_SCORE_PCT:
            st.balloons()
            st.success(f"🎊 **CHÚC MỪNG!** Bạn đã ĐẠT (PASS) — {score}/{len(test_batch)} câu đúng.")
        else:
            st.error(f"😢 **KHÔNG ĐẠT (FAIL)**. Bạn đạt {int(final_score_scaled)}/100.")
            
        if st.button("🔄 Làm lại bài thi", key=f"{test_key_prefix}_restart_btn"):
            for i, q in enumerate(test_batch, start=1):
                st.session_state.pop(f"{test_key_prefix}_q_{i}_{hash(q['question'])}", None)
            st.session_state[f"{test_key_prefix}_started"] = False
            st.session_state[f"{test_key_prefix}_submitted"] = False
            st.rerun()

def build_docwise_test_questions():
    """
    Tạo bộ đề 100 câu cho Docwise Test theo công thức:
      - 1 paragraph random từ PL3 (Bài đọc hiểu) → lấy TẤT CẢ câu hỏi của paragraph đó
      - 1 paragraph random từ PL4 (Luật & quy trình) → lấy TẤT CẢ câu hỏi của paragraph đó
      - 10 câu random từ PL5 (Chuyên ngành)
      - 30 câu random từ PL2 (Từ vựng, thuật ngữ)
      - Phần còn lại (100 - PL3_qs - PL4_qs - 10 - 30) từ PL1 (Ngữ pháp chung)
    Trả về: (test_questions, info_dict) hoặc ([], None) nếu lỗi
    """
    TOTAL = 100
    N_PL5 = 10
    N_PL2 = 30

    # --- Load các phụ lục ---
    all_pl3 = parse_pl3_passage_bank("PL3.docx")
    all_pl4 = parse_pl4_law_process("PL4.docx")
    all_pl5 = parse_pl5_specialized("PL5.docx")
    all_pl2 = parse_pl2("PL2.docx")
    all_pl1 = parse_pl1("PL1.docx")

    errors = []
    if not all_pl3: errors.append("PL3 (Bài đọc hiểu)")
    if not all_pl4: errors.append("PL4 (Luật & quy trình)")
    if not all_pl5: errors.append("PL5 (Chuyên ngành)")
    if not all_pl2: errors.append("PL2 (Từ vựng)")
    if not all_pl1: errors.append("PL1 (Ngữ pháp)")
    if errors:
        return [], {"error": f"Không đọc được file: {', '.join(errors)}"}

    # --- Chọn 1 paragraph ngẫu nhiên từ PL3 ---
    pl3_groups = {}
    for q in all_pl3:
        g = q.get('group', 'Unknown')
        pl3_groups.setdefault(g, []).append(q)
    chosen_pl3_group = random.choice(list(pl3_groups.keys()))
    qs_pl3 = pl3_groups[chosen_pl3_group]

    # --- Chọn 1 paragraph ngẫu nhiên từ PL4 ---
    pl4_groups = {}
    for q in all_pl4:
        g = q.get('group', 'Unknown')
        pl4_groups.setdefault(g, []).append(q)
    chosen_pl4_group = random.choice(list(pl4_groups.keys()))
    qs_pl4 = pl4_groups[chosen_pl4_group]

    # --- Chọn 10 câu từ PL5 ---
    qs_pl5 = random.sample(all_pl5, min(N_PL5, len(all_pl5)))

    # --- Chọn 30 câu từ PL2 ---
    qs_pl2 = random.sample(all_pl2, min(N_PL2, len(all_pl2)))

    # --- Tính số câu còn lại từ PL1 ---
    used = len(qs_pl3) + len(qs_pl4) + len(qs_pl5) + len(qs_pl2)
    n_pl1 = max(0, TOTAL - used)
    qs_pl1 = random.sample(all_pl1, min(n_pl1, len(all_pl1)))

    # --- Gắn nhãn nguồn cho mỗi câu (để phân biệt khi hiển thị) ---
    def tag(qs, source_label):
        for q in qs:
            q['_source'] = source_label
        return qs

    tag(qs_pl3, "PL3")
    tag(qs_pl4, "PL4")
    tag(qs_pl5, "PL5")
    tag(qs_pl2, "PL2")
    tag(qs_pl1, "PL1")

    # --- Ghép theo thứ tự: PL3 → PL4 → (shuffle PL5+PL2+PL1) ---
    mixed_tail = qs_pl5 + qs_pl2 + qs_pl1
    random.shuffle(mixed_tail)
    final_questions = qs_pl3 + qs_pl4 + mixed_tail

    # Đánh lại số thứ tự toàn cục
    for idx, q in enumerate(final_questions, start=1):
        q['global_number'] = idx

    info = {
        "pl3_group": chosen_pl3_group,
        "pl3_count": len(qs_pl3),
        "pl4_group": chosen_pl4_group,
        "pl4_count": len(qs_pl4),
        "pl5_count": len(qs_pl5),
        "pl2_count": len(qs_pl2),
        "pl1_count": len(qs_pl1),
        "total": len(final_questions),
    }
    return final_questions, info


def display_docwise_test_mode(bank_name, key_prefix="docwise_test"):
    """
    Màn hình Test Mode riêng cho Docwise: 100 câu từ tất cả phụ lục.
    - Trang bắt đầu: chỉ hiển thị tổng 100 câu và điểm pass 70/100
    - Trong khi làm bài: ẩn chi tiết phụ lục, paragraph đánh số 1/2 theo thứ tự xuất hiện
    """
    TOTAL_QUESTIONS = 100
    PASS_SCORE = 70          # điểm pass cố định
    bank_slug = "docwise"
    test_key_prefix = f"{key_prefix}_{bank_slug}"

    for k in [f"{test_key_prefix}_started", f"{test_key_prefix}_submitted"]:
        if k not in st.session_state:
            st.session_state[k] = False
    if f"{test_key_prefix}_questions" not in st.session_state:
        st.session_state[f"{test_key_prefix}_questions"] = []
    if f"{test_key_prefix}_info" not in st.session_state:
        st.session_state[f"{test_key_prefix}_info"] = {}

    score = 0

    # ===================== TRANG BẮT ĐẦU =====================
    if not st.session_state[f"{test_key_prefix}_started"]:
        st.markdown('<div class="result-title"><h3>📝 LÀM BÀI TEST TỔNG HỢP DOCWISE</h3></div>', unsafe_allow_html=True)
        st.markdown(f"""
        <div style="background:rgba(255,215,0,0.08); border-left:4px solid #FFD700; padding:16px 22px; border-radius:8px; margin-bottom:20px; font-size:18px; text-align:center; line-height:2;">
        📋 Tổng số câu hỏi: <b>{TOTAL_QUESTIONS} câu</b><br>
        🎯 Điểm đạt (PASS): <b>{PASS_SCORE}/{TOTAL_QUESTIONS}</b>
        </div>
        """, unsafe_allow_html=True)

        if st.button("🚀 Bắt đầu Bài Test 100 câu", key=f"{test_key_prefix}_start_btn"):
            with st.spinner("⏳ Đang tạo đề thi ngẫu nhiên..."):
                test_qs, info = build_docwise_test_questions()
            if not test_qs:
                err = info.get("error", "Lỗi không xác định") if info else "Lỗi không xác định"
                st.error(f"❌ {err}")
                return
            st.session_state[f"{test_key_prefix}_questions"] = test_qs
            st.session_state[f"{test_key_prefix}_info"] = info
            st.session_state[f"{test_key_prefix}_started"] = True
            st.session_state[f"{test_key_prefix}_submitted"] = False
            st.rerun()
        return

    # ===================== ĐANG LÀM BÀI =====================
    test_batch = st.session_state[f"{test_key_prefix}_questions"]
    current_passage_id = None
    passage_display_counter = 0   # đếm paragraph theo thứ tự xuất hiện: 1, 2, ...

    if not st.session_state[f"{test_key_prefix}_submitted"]:
        st.markdown(f'<div class="result-title"><h3>⏳ ĐANG LÀM BÀI TEST — {TOTAL_QUESTIONS} CÂU</h3></div>', unsafe_allow_html=True)

        for i, q in enumerate(test_batch, start=1):
            q_key = f"{test_key_prefix}_q_{i}_{hash(q['question'])}"
            translation_key = f"trans_{q_key}"
            is_active = (translation_key == st.session_state.active_translation_key)

            # Hiển thị đoạn văn (PL3 / PL4) — tiêu đề đổi thành "Paragraph 1", "Paragraph 2"
            passage_content = q.get('paragraph_content', '').strip()
            group_name = q.get('group', '')
            if passage_content:
                passage_id = f"passage_{group_name}_{hash(passage_content)}"
                is_passage_active = (passage_id == st.session_state.active_passage_translation)
                if passage_id != current_passage_id:
                    passage_display_counter += 1
                    display_group_label = f"Paragraph {passage_display_counter}"
                    st.markdown(f'<div class="paragraph-title">**{display_group_label}**</div>', unsafe_allow_html=True)
                    st.markdown(f'<div class="paragraph-content-box">{passage_content}</div>', unsafe_allow_html=True)
                    if not q.get('is_content_only'):
                        st.toggle("🌐 Dịch đoạn văn sang Tiếng Việt", value=is_passage_active,
                                  key=f"toggle_passage_{passage_id}",
                                  on_change=on_passage_translate_toggle, args=(passage_id,))
                        if is_passage_active:
                            translated_passage = st.session_state.passage_translations_cache.get(passage_id)
                            if not isinstance(translated_passage, str):
                                translated_passage = translate_passage_content(passage_content)
                                st.session_state.passage_translations_cache[passage_id] = translated_passage
                            st.markdown(f"""
                            <div data-testid="stAlert" class="stAlert stAlert-info">
                                <div style="font-size:18px; line-height:1.6; color:white; padding:10px;">
                                    <strong style="color:#FFD700;">[Bản dịch Đoạn văn]</strong>
                                    <div class="paragraph-content-box" style="white-space:pre-wrap; margin-bottom:0; padding:10px; background-color:rgba(0,0,0,0.5); border-left:3px solid #00d4ff;">
                                    {translated_passage}
                                    </div>
                                </div>
                            </div>""", unsafe_allow_html=True)
                    st.markdown("---")
                    current_passage_id = passage_id

            if q.get('is_content_only'):
                continue

            # Số thứ tự liên tục 1→100
            st.markdown(f'<div class="bank-question-text">{i}. {q["question"]}</div>', unsafe_allow_html=True)

            st.toggle("🌐 Dịch Câu hỏi & Đáp án sang Tiếng Việt", value=is_active,
                      key=f"toggle_{translation_key}",
                      on_change=on_translate_toggle, args=(translation_key,))
            if is_active:
                translated_content = st.session_state.translations.get(translation_key)
                if not isinstance(translated_content, str):
                    st.session_state.translations[translation_key] = translate_text(build_translation_text_for_qa(q))
                    translated_content = st.session_state.translations[translation_key]
                st.info(translated_content, icon="🌐")

            default_val = st.session_state.get(q_key, None)
            # Ẩn dấu (*) khỏi nhãn của radio button
            options_clean = [opt.replace("(*)", "").strip() for opt in q["options"]]
            st.radio("", options_clean,
                     index=options_clean.index(default_val.replace("(*)", "").strip()) if default_val and default_val.replace("(*)", "").strip() in options_clean else None,
                     key=q_key)
            st.markdown('<div class="question-separator"></div>', unsafe_allow_html=True)

        if st.button("✅ Nộp bài Test", key=f"{test_key_prefix}_submit_btn"):
            st.session_state[f"{test_key_prefix}_submitted"] = True
            st.session_state[f"{test_key_prefix}_scroll_result"] = True
            st.session_state.active_translation_key = None
            st.session_state.active_passage_translation = None
            st.rerun()

    # ===================== KẾT QUẢ =====================
    else:
        st.markdown('<div id="docwise-test-result"></div>', unsafe_allow_html=True)
        if st.session_state.pop(f"{test_key_prefix}_scroll_result", False):
            st.markdown("""
            <script>
setTimeout(function() {
    // Streamlit dùng overflow scroll trên div gốc, không phải window
    var containers = [
        document.querySelector('.main .block-container'),
        document.querySelector('[data-testid="stAppViewContainer"] > section'),
        document.querySelector('.main'),
        document.documentElement,
        document.body
    ];
    for (var c of containers) {
        if (c) { c.scrollTop = c.scrollHeight; }
    }
    window.scrollTo(0, document.body.scrollHeight);
    window.parent.scrollTo(0, window.parent.document.body.scrollHeight);
}, 400);
</script>
            """, unsafe_allow_html=True)
        st.markdown('<div class="result-title"><h3>🎉 KẾT QUẢ BÀI TEST TỔNG HỢP</h3></div>', unsafe_allow_html=True)
        current_passage_id = None
        passage_display_counter = 0

        for i, q in enumerate(test_batch, start=1):
            q_key = f"{test_key_prefix}_q_{i}_{hash(q['question'])}"
            selected_opt = st.session_state.get(q_key)
            correct = clean_text(q["answer"])
            is_correct = (selected_opt is not None) and (clean_text(selected_opt) == correct)
            translation_key = f"trans_{q_key}"
            is_active = (translation_key == st.session_state.active_translation_key)

            # Hiển thị đoạn văn — tiêu đề "Paragraph 1", "Paragraph 2"
            passage_content = q.get('paragraph_content', '').strip()
            group_name = q.get('group', '')
            if passage_content:
                passage_id = f"passage_{group_name}_{hash(passage_content)}"
                is_passage_active = (passage_id == st.session_state.active_passage_translation)
                if passage_id != current_passage_id:
                    passage_display_counter += 1
                    display_group_label = f"Paragraph {passage_display_counter}"
                    st.markdown(f'<div class="paragraph-title">**{display_group_label}**</div>', unsafe_allow_html=True)
                    st.markdown(f'<div class="paragraph-content-box">{passage_content}</div>', unsafe_allow_html=True)
                    if not q.get('is_content_only'):
                        st.toggle("🌐 Dịch đoạn văn sang Tiếng Việt", value=is_passage_active,
                                  key=f"toggle_passage_{passage_id}",
                                  on_change=on_passage_translate_toggle, args=(passage_id,))
                        if is_passage_active:
                            translated_passage = st.session_state.passage_translations_cache.get(passage_id)
                            if not isinstance(translated_passage, str):
                                translated_passage = translate_passage_content(passage_content)
                                st.session_state.passage_translations_cache[passage_id] = translated_passage
                            st.markdown(f"""
                            <div data-testid="stAlert" class="stAlert stAlert-info">
                                <div style="font-size:18px; line-height:1.6; color:white; padding:10px;">
                                    <strong style="color:#FFD700;">[Bản dịch Đoạn văn]</strong>
                                    <div class="paragraph-content-box" style="white-space:pre-wrap; margin-bottom:0; padding:10px; background-color:rgba(0,0,0,0.5); border-left:3px solid #00d4ff;">
                                    {translated_passage}
                                    </div>
                                </div>
                            </div>""", unsafe_allow_html=True)
                    st.markdown("---")
                    current_passage_id = passage_id

            if q.get('is_content_only'):
                continue

            # Câu hỏi và đáp án
            st.markdown(f'<div class="bank-question-text">{i}. {q["question"]}</div>', unsafe_allow_html=True)
            st.toggle("🌐 Dịch Câu hỏi & Đáp án sang Tiếng Việt", value=is_active,
                      key=f"toggle_{translation_key}",
                      on_change=on_translate_toggle, args=(translation_key,))
            if is_active:
                translated_content = st.session_state.translations.get(translation_key)
                if not isinstance(translated_content, str):
                    st.session_state.translations[translation_key] = translate_text(build_translation_text_for_qa(q))
                    translated_content = st.session_state.translations[translation_key]
                st.info(translated_content, icon="🌐")

            for opt in q["options"]:
                opt_clean = clean_text(opt)
                # Ẩn dấu (*) nếu chưa nộp bài trong Test Mode
                opt_display = opt.replace("(*)", "").strip()
                if opt_clean == correct:
                    ans_class = "bank-answer-text answer-correct"
                elif selected_opt is not None and clean_text(selected_opt) == opt_clean:
                    ans_class = "bank-answer-text answer-selected-wrong"
                else:
                    ans_class = "bank-answer-text answer-wrong"
                st.markdown(f'<div class="{ans_class}">{opt_display}</div>', unsafe_allow_html=True)

            if is_correct:
                score += 1
            st.markdown('<div class="question-separator"></div>', unsafe_allow_html=True)

        total_q = len(test_batch)
        st.markdown(f'<div class="result-title"><h3>🎯 KẾT QUẢ: {score}/{total_q}</h3></div>', unsafe_allow_html=True)
        if score >= PASS_SCORE:
            st.balloons()
            st.success(f"🎊 **CHÚC MỪNG!** Bạn đã ĐẠT (PASS) — {score}/{total_q} câu đúng.")
        else:
            st.error(f"😢 **KHÔNG ĐẠT (FAIL)**. Cần {PASS_SCORE} câu đúng để Đạt. Bạn đạt {score}/{total_q}.")

        if st.button("🔄 Làm lại Bài Test (Đề mới)", key=f"{test_key_prefix}_restart_btn"):
            for i, q in enumerate(test_batch, start=1):
                st.session_state.pop(f"{test_key_prefix}_q_{i}_{hash(q['question'])}", None)
            st.session_state.pop(f"{test_key_prefix}_questions", None)
            st.session_state.pop(f"{test_key_prefix}_info", None)
            st.session_state[f"{test_key_prefix}_started"] = False
            st.session_state[f"{test_key_prefix}_submitted"] = False
            st.rerun()




def display_test_mode(questions, bank_name, key_prefix="test"):
    TOTAL_QUESTIONS = 50
    PASS_SCORE = math.ceil(TOTAL_QUESTIONS * 0.75)  # 38/50
    bank_slug = bank_name.split()[-1].lower()
    test_key_prefix = f"{key_prefix}_{bank_slug}"

    if f"{test_key_prefix}_started" not in st.session_state:
        st.session_state[f"{test_key_prefix}_started"] = False
    if f"{test_key_prefix}_submitted" not in st.session_state:
        st.session_state[f"{test_key_prefix}_submitted"] = False
    if f"{test_key_prefix}_questions" not in st.session_state:
        st.session_state[f"{test_key_prefix}_questions"] = []

    score = 0

    if not st.session_state[f"{test_key_prefix}_started"]:
        st.markdown('<div class="result-title"><h3>📝 LÀM BÀI TEST 50 CÂU</h3></div>', unsafe_allow_html=True)
        st.markdown(f"""
        <div style="background:rgba(255,215,0,0.08); border-left:4px solid #FFD700; padding:16px 22px; border-radius:8px; margin-bottom:20px; font-size:18px; text-align:center; line-height:2;">
        📋 Tổng số câu hỏi: <b>{TOTAL_QUESTIONS} câu</b><br>
        🎯 Điểm đạt (PASS): <b>{PASS_SCORE}/{TOTAL_QUESTIONS}</b>
        </div>
        """, unsafe_allow_html=True)

        if st.button("🚀 Bắt đầu Bài Test", key=f"{test_key_prefix}_start_btn"):
            st.session_state[f"{test_key_prefix}_questions"] = get_random_questions(questions, TOTAL_QUESTIONS)
            st.session_state[f"{test_key_prefix}_started"] = True
            st.session_state[f"{test_key_prefix}_submitted"] = False
            st.session_state.current_mode = "test"
            st.rerun()
        return

    # Logic hiển thị đoạn văn trong Test Mode (chỉ hiển thị 1 lần cho mỗi đoạn)
    test_batch = st.session_state[f"{test_key_prefix}_questions"]
    current_passage_id = None

    if not st.session_state[f"{test_key_prefix}_submitted"]:
        st.markdown('<div class="result-title"><h3>⏳ ĐANG LÀM BÀI TEST</h3></div>', unsafe_allow_html=True)
        for i, q in enumerate(test_batch, start=1):
            q_key = f"{test_key_prefix}_q_{i}_{hash(q['question'])}" 
            translation_key = f"trans_{q_key}"
            is_active = (translation_key == st.session_state.active_translation_key)
            
            # --- BỔ SUNG: HIỂN THỊ ĐOẠN VĂN (CHO PL3) ---
            passage_content = q.get('paragraph_content', '').strip()
            group_name = q.get('group', '')
            
            if passage_content:
                passage_id = f"passage_{group_name}_{hash(passage_content)}"
                is_passage_active = (passage_id == st.session_state.active_passage_translation)

                if passage_id != current_passage_id:
                     # 1. In đậm, đổi màu tiêu đề
                    st.markdown(f'<div class="paragraph-title">**{group_name}**</div>', unsafe_allow_html=True) 
                    
                    # 2. Hiển thị nội dung đoạn văn gốc
                    st.markdown(f'<div class="paragraph-content-box">{passage_content}</div>', unsafe_allow_html=True)
                    
                    # 3. Thêm Nút Dịch Đoạn Văn
                    st.toggle(
                        "🌐 Dịch đoạn văn sang Tiếng Việt", 
                        value=is_passage_active, 
                        key=f"toggle_passage_{passage_id}",
                        on_change=on_passage_translate_toggle,
                        args=(passage_id,)
                    )
                    
                    # 4. Hiển thị Bản Dịch Đoạn Văn
                    if is_passage_active:
                        translated_passage = st.session_state.passage_translations_cache.get(passage_id)
                        if not isinstance(translated_passage, str):
                            # GỌI HÀM DỊCH CHỈ ĐOẠAN VĂN
                            translated_passage = translate_passage_content(passage_content)
                            st.session_state.passage_translations_cache[passage_id] = translated_passage

                        st.markdown(f"""
                        <div data-testid="stAlert" class="stAlert stAlert-info">
                            <div style="font-size: 18px; line-height: 1.6; color: white; padding: 10px;">
                                <strong style="color: #FFD700;">[Bản dịch Đoạn văn]</strong>
                                <div class="paragraph-content-box" style="white-space: pre-wrap; margin-bottom: 0px; padding: 10px; background-color: rgba(0, 0, 0, 0.5); border-left: 3px solid #00d4ff;">
                                {translated_passage}
                                </div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)

                    st.markdown("---")
                    current_passage_id = passage_id
            # --- KẾT THÚC BỔ SUNG ---

            # Hiển thị câu hỏi (SỬ DỤNG SỐ THỨ TỰ CỤC BỘ NẾU LÀ PL3, NẾU KHÔNG DÙNG SỐ THỨ TỰ TOÀN CỤC)
            if q.get('group', '').startswith('Paragraph'):
                # Dùng số thứ tự cục bộ (number) nếu là bài đọc hiểu
                display_num = q.get('number', i) 
            else:
                # Dùng số thứ tự toàn cục (i) cho các ngân hàng khác
                display_num = i
            st.markdown(f'<div class="bank-question-text">{display_num}. {q["question"]}</div>', unsafe_allow_html=True)

            # Nút Dịch Q&A ở dưới
            st.toggle(
                "🌐 Dịch Câu hỏi & Đáp án sang Tiếng Việt", 
                value=is_active, 
                key=f"toggle_{translation_key}",
                on_change=on_translate_toggle,
                args=(translation_key,)
            )

            # Hiển thị Bản Dịch Q&A
            if is_active:
                translated_content = st.session_state.translations.get(translation_key)
                
                if not isinstance(translated_content, str):
                    # GỌI HÀM MỚI ĐỂ GỬI CHỈ CÂU HỎI VÀ ĐÁP ÁN ĐI DỊCH
                    full_text_to_translate = build_translation_text_for_qa(q)
                    st.session_state.translations[translation_key] = translate_text(full_text_to_translate)
                    translated_content = st.session_state.translations[translation_key]

                st.info(translated_content, icon="🌐")

            # Hiển thị Radio Button (không chọn mặc định)
            default_val = st.session_state.get(q_key, None)
            # Ẩn dấu (*) khỏi nhãn của radio button
            options_clean = [opt.replace("(*)", "").strip() for opt in q["options"]]
            st.radio("", options_clean, index=options_clean.index(default_val.replace("(*)", "").strip()) if default_val and default_val.replace("(*)", "").strip() in options_clean else None, key=q_key)
            st.markdown('<div class="question-separator"></div>', unsafe_allow_html=True)
            
        if st.button("✅ Nộp bài Test", key=f"{test_key_prefix}_submit_btn"):
            st.session_state[f"{test_key_prefix}_submitted"] = True
            st.session_state[f"{test_key_prefix}_scroll_result"] = True
            st.session_state.active_translation_key = None # Tắt dịch Q&A khi nộp
            st.session_state.active_passage_translation = None # Tắt dịch Passage khi nộp
            st.rerun()
            
    else:
        st.markdown('<div id="regular-test-result"></div>', unsafe_allow_html=True)
        if st.session_state.pop(f"{test_key_prefix}_scroll_result", False):
            st.markdown("""
            <script>
setTimeout(function() {
    // Streamlit dùng overflow scroll trên div gốc, không phải window
    var containers = [
        document.querySelector('.main .block-container'),
        document.querySelector('[data-testid="stAppViewContainer"] > section'),
        document.querySelector('.main'),
        document.documentElement,
        document.body
    ];
    for (var c of containers) {
        if (c) { c.scrollTop = c.scrollHeight; }
    }
    window.scrollTo(0, document.body.scrollHeight);
    window.parent.scrollTo(0, window.parent.document.body.scrollHeight);
}, 400);
</script>
            """, unsafe_allow_html=True)
        st.markdown('<div class="result-title"><h3>🎉 KẾT QUẢ BÀI TEST</h3></div>', unsafe_allow_html=True)
        
        for i, q in enumerate(test_batch, start=1):
            q_key = f"{test_key_prefix}_q_{i}_{hash(q['question'])}" 
            selected_opt = st.session_state.get(q_key)
            correct = clean_text(q["answer"])
            is_correct = (selected_opt is not None) and (clean_text(selected_opt) == correct)
            translation_key = f"trans_{q_key}"
            is_active = (translation_key == st.session_state.active_translation_key)

            # --- BỔ SUNG: HIỂN THỊ ĐOẠN VĂN (CHO PL3) ---
            passage_content = q.get('paragraph_content', '').strip()
            group_name = q.get('group', '')
            
            if passage_content:
                passage_id = f"passage_{group_name}_{hash(passage_content)}"
                is_passage_active = (passage_id == st.session_state.active_passage_translation)

                if passage_id != current_passage_id:
                     # 1. In đậm, đổi màu tiêu đề
                    st.markdown(f'<div class="paragraph-title">**{group_name}**</div>', unsafe_allow_html=True) 
                    
                    # 2. Hiển thị nội dung đoạn văn gốc
                    st.markdown(f'<div class="paragraph-content-box">{passage_content}</div>', unsafe_allow_html=True)
                    
                    # 3. Thêm Nút Dịch Đoạn Văn
                    st.toggle(
                        "🌐 Dịch đoạn văn sang Tiếng Việt", 
                        value=is_passage_active, 
                        key=f"toggle_passage_{passage_id}",
                        on_change=on_passage_translate_toggle,
                        args=(passage_id,)
                    )
                    
                    # 4. Hiển thị Bản Dịch Đoạn Văn
                    if is_passage_active:
                        translated_passage = st.session_state.passage_translations_cache.get(passage_id)
                        if not isinstance(translated_passage, str):
                            # GỌI HÀM DỊCH CHỈ ĐOẠN VĂN
                            translated_passage = translate_passage_content(passage_content)
                            st.session_state.passage_translations_cache[passage_id] = translated_passage

                        st.markdown(f"""
                        <div data-testid="stAlert" class="stAlert stAlert-info">
                            <div style="font-size: 18px; line-height: 1.6; color: white; padding: 10px;">
                                <strong style="color: #FFD700;">[Bản dịch Đoạn văn]</strong>
                                <div class="paragraph-content-box" style="white-space: pre-wrap; margin-bottom: 0px; padding: 10px; background-color: rgba(0, 0, 0, 0.5); border-left: 3px solid #00d4ff;">
                                {translated_passage}
                                </div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)

                    st.markdown("---")
                    current_passage_id = passage_id
            # --- KẾT THÚC BỔ SUNG ---

            # Hiển thị câu hỏi (SỬ DỤNG SỐ THỨ TỰ CỤC BỘ NẾU LÀ PL3, NẾU KHÔNG DÙNG SỐ THỨ TỰ TOÀN CỤC)
            if q.get('group', '').startswith('Paragraph'):
                # Dùng số thứ tự cục bộ (number) nếu là bài đọc hiểu
                display_num = q.get('number', i) 
            else:
                # Dùng số thứ tự toàn cục (i) cho các ngân hàng khác
                display_num = i
            st.markdown(f'<div class="bank-question-text">{display_num}. {q["question"]}</div>', unsafe_allow_html=True)

            # Nút Dịch Q&A ở dưới
            st.toggle(
                "🌐 Dịch Câu hỏi & Đáp án sang Tiếng Việt", 
                value=is_active, 
                key=f"toggle_{translation_key}",
                on_change=on_translate_toggle,
                args=(translation_key,)
            )

            # Hiển thị Bản Dịch Q&A
            if is_active:
                translated_content = st.session_state.translations.get(translation_key)
                
                if not isinstance(translated_content, str):
                    # GỌI HÀM MỚI ĐỂ GỬI CHỈ CÂU HỎI VÀ ĐÁP ÁN ĐI DỊCH
                    full_text_to_translate = build_translation_text_for_qa(q)
                    st.session_state.translations[translation_key] = translate_text(full_text_to_translate)
                    translated_content = st.session_state.translations[translation_key]

                st.info(translated_content, icon="🌐")
            
            # Hiển thị Đáp án (KẾT QUẢ)
            for opt in q["options"]:
                opt_clean = clean_text(opt)
                # Ẩn dấu (*) nếu chưa nộp bài
                opt_display = opt.replace("(*)", "").strip()

                if opt_clean == correct:
                    ans_class = "bank-answer-text answer-correct"
                else:
                    ans_class = "bank-answer-text answer-wrong"
                    
                st.markdown(f'<div class="{ans_class}">{opt_display}</div>', unsafe_allow_html=True)

            if is_correct: score += 1
            st.markdown('<div class="question-separator"></div>', unsafe_allow_html=True)
        
        total_q = len(test_batch)
        st.markdown(f'<div class="result-title"><h3>🎯 KẾT QUẢ: {score}/{total_q}</h3></div>', unsafe_allow_html=True)

        if score >= PASS_SCORE:
            st.balloons()
            st.success(f"🎊 **CHÚC MỪNG!** Bạn đã ĐẠT (PASS) — {score}/{total_q} câu đúng.")
        else:
            st.error(f"😢 **KHÔNG ĐẠT (FAIL)**. Cần {PASS_SCORE} câu đúng để Đạt. Bạn đạt {score}/{total_q}.")

        if st.button("🔄 Làm lại Bài Test", key=f"{test_key_prefix}_restart_btn"):
            for i, q in enumerate(test_batch, start=1):
                st.session_state.pop(f"{test_key_prefix}_q_{i}_{hash(q['question'])}", None)
            st.session_state.pop(f"{test_key_prefix}_questions", None)
            st.session_state[f"{test_key_prefix}_started"] = False
            st.session_state[f"{test_key_prefix}_submitted"] = False
            st.rerun()

# ====================================================
# 🖥️ GIAO DIỆN STREAMLIT
# ====================================================
def build_appendix_specific_test_questions(appendix_short):
    """
    Tạo bộ đề cho từng phụ lục:
    - PL1, PL2, PL5: 50 câu ngẫu nhiên.
    - PL3: 3 paragraph ngẫu nhiên từ PL3.
    - PL4: 3 paragraph ngẫu nhiên từ PL4.
    """
    if any(x in appendix_short for x in ["Phụ lục 1", "Phụ lục 2", "Phụ lục 5"]):
        if "Phụ lục 1" in appendix_short: qs = parse_pl1("PL1.docx")
        elif "Phụ lục 2" in appendix_short: qs = parse_pl2("PL2.docx")
        else: qs = parse_pl5_specialized("PL5.docx")
        
        if not qs: return [], 0
        return random.sample(qs, min(50, len(qs))), 50
    elif "Phụ lục 3" in appendix_short:
        all_pl3 = parse_pl3_passage_bank("PL3.docx")
        groups = {}
        for q in all_pl3:
            g = q.get('group', 'Unknown')
            groups.setdefault(g, []).append(q)
        if not groups: return [], 0
        chosen = random.sample(list(groups.keys()), min(3, len(groups)))
        final = []
        for c in chosen:
            final.extend(groups[c])
        return final, len(final)
    elif "Phụ lục 4" in appendix_short:
        all_pl4 = parse_pl4_law_process("PL4.docx")
        groups = {}
        for q in all_pl4:
            g = q.get('group', 'Unknown')
            groups.setdefault(g, []).append(q)
        if not groups: return [], 0
        chosen = random.sample(list(groups.keys()), min(3, len(groups)))
        final = []
        for c in chosen:
            final.extend(groups[c])
        return final, len(final)
    return [], 0

def display_appendix_test_mode(appendix_full_name):
    appendix_short = appendix_full_name.split(':')[0].strip()
    test_key_prefix = f"appendix_test_{appendix_short.replace(' ', '_').lower()}"
    PASS_THRESHOLD = 70 # điểm pass 70/100
    
    if f"{test_key_prefix}_started" not in st.session_state:
        st.session_state[f"{test_key_prefix}_started"] = False
    if f"{test_key_prefix}_submitted" not in st.session_state:
        st.session_state[f"{test_key_prefix}_submitted"] = False
    if f"{test_key_prefix}_questions" not in st.session_state:
        st.session_state[f"{test_key_prefix}_questions"] = []

    if not st.session_state[f"{test_key_prefix}_started"]:
        st.markdown(f'<div class="result-title"><h3>📝 LÀM BÀI TEST {appendix_short.upper()}</h3></div>', unsafe_allow_html=True)
        
        # Mô tả cơ cấu đề
        if any(x in appendix_short for x in ["Phụ lục 1", "Phụ lục 2", "Phụ lục 5"]):
            desc = "📋 Tổng số câu hỏi: <b>50 câu (ngẫu nhiên)</b>"
        else:
            desc = f"📋 Cơ cấu đề: <b>3 Paragraph ngẫu nhiên (từ {appendix_short})</b>"
            
        st.markdown(f"""
        <div style="background:rgba(255,215,0,0.08); border-left:4px solid #FFD700; padding:16px 22px; border-radius:8px; margin-bottom:20px; font-size:18px; text-align:center; line-height:2;">
        {desc}<br>
        🎯 Điểm đạt (PASS): <b>70/100</b>
        </div>
        """, unsafe_allow_html=True)

        if st.button(f"🚀 Bắt đầu Bài Test {appendix_short}", key=f"{test_key_prefix}_start_btn"):
            with st.spinner("⏳ Đang tạo đề thi..."):
                test_qs, total_val = build_appendix_specific_test_questions(appendix_short)
            if not test_qs:
                st.error("❌ Không thể tạo đề thi. Vui lòng kiểm tra lại file dữ liệu.")
                return
            st.session_state[f"{test_key_prefix}_questions"] = test_qs
            st.session_state[f"{test_key_prefix}_started"] = True
            st.session_state[f"{test_key_prefix}_submitted"] = False
            st.rerun()
        return

    test_batch = st.session_state[f"{test_key_prefix}_questions"]
    current_passage_id = None
    passage_counter = 0

    if not st.session_state[f"{test_key_prefix}_submitted"]:
        st.markdown(f'<div class="result-title"><h3>⏳ ĐANG LÀM BÀI TEST {appendix_short.upper()}</h3></div>', unsafe_allow_html=True)
        for i, q in enumerate(test_batch, start=1):
            q_key = f"{test_key_prefix}_q_{i}_{hash(q['question'])}"
            translation_key = f"trans_{q_key}"
            is_active = (translation_key == st.session_state.active_translation_key)
            
            # Đoạn văn
            passage_content = q.get('paragraph_content', '').strip()
            group_name = q.get('group', '')
            if passage_content:
                passage_id = f"passage_{group_name}_{hash(passage_content)}"
                is_passage_active = (passage_id == st.session_state.active_passage_translation)
                if passage_id != current_passage_id:
                    passage_counter += 1
                    st.markdown(f'<div class="paragraph-title">**Paragraph {passage_counter}**</div>', unsafe_allow_html=True)
                    st.markdown(f'<div class="paragraph-content-box">{passage_content}</div>', unsafe_allow_html=True)
                    if not q.get('is_content_only'):
                        st.toggle("🌐 Dịch đoạn văn", value=is_passage_active, key=f"toggle_passage_{passage_id}_{test_key_prefix}", on_change=on_passage_translate_toggle, args=(passage_id,))
                        if is_passage_active:
                            trans = st.session_state.passage_translations_cache.get(passage_id)
                            if not trans:
                                trans = translate_passage_content(passage_content)
                                st.session_state.passage_translations_cache[passage_id] = trans
                            st.info(trans)
                    current_passage_id = passage_id

            if q.get('is_content_only'):
                continue

            st.markdown(f'<div class="bank-question-text">{i}. {q["question"]}</div>', unsafe_allow_html=True)
            st.toggle("🌐 Dịch Q&A", value=is_active, key=f"toggle_{translation_key}_{test_key_prefix}", on_change=on_translate_toggle, args=(translation_key,))
            if is_active:
                trans_qa = st.session_state.translations.get(translation_key)
                if not isinstance(trans_qa, str):
                    st.session_state.translations[translation_key] = translate_text(build_translation_text_for_qa(q))
                    trans_qa = st.session_state.translations[translation_key]
                st.info(trans_qa)

            options_clean = [opt.replace("(*)", "").strip() for opt in q["options"]]
            st.radio("", options_clean, index=None, key=q_key)
            st.markdown('<div class="question-separator"></div>', unsafe_allow_html=True)
            
        if st.button("✅ Nộp bài", key=f"{test_key_prefix}_submit_btn"):
            st.session_state[f"{test_key_prefix}_submitted"] = True
            st.rerun()
    else:
        score = 0
        current_passage_id = None
        passage_counter = 0
        for i, q in enumerate(test_batch, start=1):
            q_key = f"{test_key_prefix}_q_{i}_{hash(q['question'])}"
            selected_opt = st.session_state.get(q_key)
            correct = clean_text(q["answer"])
            is_correct = (selected_opt is not None) and (clean_text(selected_opt) == correct)
            
            # Rendering results
            passage_content = q.get('paragraph_content', '').strip()
            group_name = q.get('group', '')
            if passage_content:
                passage_id = f"passage_{group_name}_{hash(passage_content)}"
                if passage_id != current_passage_id:
                    passage_counter += 1
                    st.markdown(f'<div class="paragraph-title">**Paragraph {passage_counter}**</div>', unsafe_allow_html=True)
                    st.markdown(f'<div class="paragraph-content-box">{passage_content}</div>', unsafe_allow_html=True)
                    current_passage_id = passage_id

            st.markdown(f'<div class="bank-question-text">{i}. {q["question"]}</div>', unsafe_allow_html=True)
            for opt in q["options"]:
                opt_clean = clean_text(opt)
                opt_display = opt.replace("(*)", "").strip()
                if opt_clean == correct:
                    ans_class = "bank-answer-text answer-correct"
                elif selected_opt is not None and clean_text(selected_opt) == opt_clean:
                    ans_class = "bank-answer-text answer-selected-wrong"
                else:
                    ans_class = "bank-answer-text answer-wrong"
                st.markdown(f'<div class="{ans_class}">{opt_display}</div>', unsafe_allow_html=True)

            if is_correct: score += 1
            st.markdown('<div class="question-separator"></div>', unsafe_allow_html=True)

        final_score = (score / len(test_batch)) * 100 if test_batch else 0
        st.markdown(f'<div class="result-title"><h3>🎯 KẾT QUẢ: {int(final_score)}/100</h3></div>', unsafe_allow_html=True)
        if final_score >= PASS_THRESHOLD:
            st.balloons()
            st.success(f"🎊 **CHÚC MỪNG!** Bạn đã ĐẠT (PASS) — {score}/{len(test_batch)} câu đúng.")
        else:
            st.error(f"😢 **KHÔNG ĐẠT (FAIL)**. Bạn đạt {int(final_score)}/100.")
        
        if st.button("🔄 Làm lại Bài Test", key=f"{test_key_prefix}_restart"):
            for i, q in enumerate(test_batch, start=1):
                st.session_state.pop(f"{test_key_prefix}_q_{i}_{hash(q['question'])}", None)
            st.session_state[f"{test_key_prefix}_started"] = False
            st.session_state[f"{test_key_prefix}_submitted"] = False
            st.rerun()

st.set_page_config(page_title="Ngân hàng trắc nghiệm", layout="wide")

PC_IMAGE_FILE = "PC2.jpg"
MOBILE_IMAGE_FILE = "mobile2.jpg"
LOGO_IMAGE_FILE = "logo.jpg"
LOGO2_IMAGE_FILE = "logo2.png"
img_pc_base64 = get_base64_encoded_file(PC_IMAGE_FILE)
img_mobile_base64 = get_base64_encoded_file(MOBILE_IMAGE_FILE)
img_logo_base64 = get_base64_encoded_file(LOGO_IMAGE_FILE)
img_logo2_base64 = get_base64_encoded_file(LOGO2_IMAGE_FILE)

# === CSS CẬP NHẬT CHO ĐOẠN VĂN (PARAGRAPH) ===
css_style = f"""
<style>
/* Đã thống nhất font nội dung là Oswald, tiêu đề là Playfair Display */
@import url('https://fonts.googleapis.com/css2?family=Rye&display=swap');

@keyframes colorShift {{
    0% {{ background-position: 0% 50%; }}
    50% {{ background-position: 100% 50%; }}
    100% {{ background-position: 0% 50%; }}
}}
@keyframes scrollRight {{
    0% {{ transform: translateX(100%); }}
    100% {{ transform: translateX(-100%); }}
}}

html, body, .stApp {{
    height: 100% !important;
    min-height: 100vh !important;
    margin: 0 !important;
    padding: 0 !important;
    overflow: auto;
    position: relative;
    font-family: 'Arial', 'Helvetica', sans-serif !important;
    color: #FFFFE0 !important;
}}

/* BACKGROUND */
.stApp {{
    background: none !important;
}}

.stApp::before {{
    content: "";
    position: fixed;
    top: 0;
    left: 0;
    width: 100%;
    height: 100%;
    background: url("data:image/jpeg;base64,{img_pc_base64}") no-repeat center top fixed;
    background-size: cover;
    filter: sepia(0.5) brightness(0.9) blur(0px);
    z-index: -1; 
    pointer-events: none;
}}

@media (max-width: 767px) {{
    .stApp::before {{
        background: url("data:image/jpeg;base64,{img_mobile_base64}") no-repeat center top scroll;
        background-size: cover;
    }}
}}

/* Nội dung nổi lên trên nền */
[data-testid="stAppViewContainer"],
[data-testid="stMainBlock"],
.main {{
    background-color: transparent !important;
}}

/* Ẩn UI */
#MainMenu, footer, header {{visibility: hidden; height: 0;}}
[data-testid="stHeader"] {{display: none;}}

/* TITLE CHÍNH */
#main-title-container {{
    position: relative; left: 0; top: 0; width: 100%;
    height: 120px; overflow: hidden;
    pointer-events: none;
    background-color: transparent; padding-top: 20px; z-index: 1200; 
    display: flex;
    align-items: center;
    justify-content: center;
}}

/* LOGO LEFT */
#logo-container {{
    position: fixed;
    top: 20px;
    left: 20px;
    z-index: 2000;
    pointer-events: none;
}}

/* Wrapper viền sáng chạy vòng quanh logo trái */
#logo-wrap {{
    position: relative;
    border-radius: 16px;
    padding: 3px;
    display: inline-block;
    overflow: hidden;
}}

@property --logo-angle {{
    syntax: '<angle>';
    initial-value: 0deg;
    inherits: false;
}}

@keyframes logo-spin-light {{
    to {{ --logo-angle: 360deg; }}
}}

#logo-wrap::before {{
    content: '';
    position: absolute;
    inset: -60%;
    background: conic-gradient(
        from var(--logo-angle, 0deg),
        transparent 0deg,
        transparent 40deg,
        #b8860b 60deg,
        #ffd700 80deg,
        #fffacd 90deg,
        #ffd700 100deg,
        #b8860b 120deg,
        transparent 140deg,
        transparent 360deg
    );
    animation: logo-spin-light 3s linear infinite;
    z-index: 0;
}}

#logo-wrap::after {{
    content: '';
    position: absolute;
    inset: 3px;
    border-radius: 13px;
    background: rgba(0,0,0,0.45);
    z-index: 1;
}}

#logo-wrap img {{
    position: relative;
    z-index: 2;
    height: 110px;
    width: auto;
    object-fit: contain;
    border-radius: 12px;
    display: block;
    filter: drop-shadow(0 2px 8px rgba(0,0,0,0.6));
}}

/* LOGO RIGHT (logo2) */
#logo2-container {{
    position: fixed;
    top: 20px;
    right: 20px;
    z-index: 2000;
    pointer-events: none;
}}

.logo2-wrap {{
    position: relative;
    display: inline-block;
    padding: 0;
}}

.logo2-wrap svg.ellipse-border {{
    position: absolute;
    inset: 0;
    width: 100%;
    height: 100%;
    z-index: 3;
    pointer-events: none;
    overflow: visible;
}}

.logo2-wrap img {{
    position: relative;
    z-index: 2;
    height: 110px;
    width: auto;
    object-fit: contain;
    display: block;
}}

@media (max-width: 767px) {{
    #logo-container {{
        top: 12px;
        left: 8px;
    }}
    #logo-wrap img {{
        height: 44px;
    }}
    #logo-wrap::after {{ inset: 2px; border-radius: 10px; }}
    #logo2-container {{
        top: 12px;
        right: 8px;
    }}
    .logo2-wrap img {{
        height: 44px;
    }}
}}

/* === BỔ SUNG CSS CHO ĐOẠN VĂN (PL3) === */

/* RESULT TITLE */
.result-title {{
    margin-top: 30px;
    margin-bottom: 30px;
    text-align: center;
}}
.result-title h3 {{
    font-family: 'Arial', 'Helvetica', sans-serif;
    font-size: 1.8rem;
    font-weight: 900;
    letter-spacing: 2px;
    color: #D4A843;
    text-shadow: 0 0 15px #D4A843, 0 0 30px rgba(212,168,67,0.8);
}}

/* Đẩy nội dung chính xuống để không bị che bởi logo */
.main > div:first-child {{
    padding-top: 130px !important; padding-bottom: 2rem !important;
}}

/* Tiêu đề Paragraph X . (In đậm, màu cam) */
.paragraph-title {{
    font-family: 'Arial', 'Helvetica', sans-serif;
    font-size: 1.4rem;
    font-weight: 700;
    color: #D4A843;
    text-shadow: 0 0 8px rgba(212, 168, 67, 0.5);
    margin-top: 20px;
    margin-bottom: 10px;
    padding: 5px 15px;
    background-color: rgba(30, 30, 30, 0.8);
    border-radius: 8px;
    display: inline-block;
    letter-spacing: 2px;
}}

/* Nội dung đoạn văn (Giữ nguyên bố cục xuống dòng) */
.paragraph-content-box {{
    /* Dùng 'white-space: pre-wrap' để giữ nguyên khoảng trắng và ngắt dòng */
    white-space: pre-wrap; 
    font-family: 'Arial', 'Helvetica', sans-serif !important;
    font-size: 20px !important; 
    line-height: 1.6;
    color: #F0F0F0; /* Màu trắng nhạt */
    padding: 15px;
    background-color: rgba(0, 0, 0, 0.7);
    border-radius: 8px;
    margin-bottom: 20px;
    border-left: 3px solid #FFA500;
    box-shadow: 0 2px 10px rgba(0, 0, 0, 0.5);
}}


/* STYLE CÂU HỎI - PC (NỀN ĐEN BAO VỪA CHỮ) */
.bank-question-text {{
    color: #FF8C00 !important;
    font-weight: 900 !important;
    letter-spacing: 0.5px !important;
    font-size: 22px !important; 
    font-family: 'Arial', 'Helvetica', sans-serif !important;
    text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.8);
    padding: 8px 15px;
    margin-bottom: 10px;
    line-height: 1.4 !important;
    background-color: rgba(0, 0, 0, 0.75);
    border-radius: 8px;
    display: inline-block; /* BAO VỪA CHỮ */
    max-width: 100%;
    box-shadow: 0 4px 10px rgba(0, 0, 0, 0.5);
}}

/* STYLE ĐÁP ÁN - PC (TRẮNG ĐẬM HƠN) */
.bank-answer-text {{
    font-family: 'Arial', 'Helvetica', sans-serif !important;
    font-weight: 700 !important;
    font-size: 22px !important; 
    padding: 5px 15px;
    margin: 2px 0;
    line-height: 1.5 !important; 
    display: block;
    color: #FFFFFF !important;
    text-shadow: 1px 1px 3px rgba(0, 0, 0, 0.9);
}}

/* RADIO BUTTONS (CHỌN ĐÁP ÁN) */
.stRadio label {{
    color: #FFFFFF !important;
    font-size: 20px !important; 
    font-weight: 700 !important;
    font-family: 'Arial', 'Helvetica', sans-serif !important;
    padding: 2px 12px;
    text-shadow: 1px 1px 3px rgba(0, 0, 0, 0.9) !important;
    background-color: transparent !important;
    border: none !important;
    display: block !important;
    margin: 4px 0 !important;
    letter-spacing: 0.3px !important;
}}

.stRadio label:hover {{
    text-shadow: 1px 1px 3px rgba(0, 0, 0, 0.9) !important;
}}

.stRadio label span, 
.stRadio label p,
.stRadio label div {{
    color: #FFFFFF !important;
    font-family: 'Arial', 'Helvetica', sans-serif !important;
    text-shadow: 1px 1px 3px rgba(0, 0, 0, 0.9) !important;
    letter-spacing: 0.3px !important;
}}

div[data-testid="stMarkdownContainer"] p {{
    font-size: 22px !important;
    font-family: 'Arial', 'Helvetica', sans-serif !important;
    color: #FFFFE0 !important;
}}

/* TIÊU ĐỀ CHÍNH - override mọi rule khác */
#bank-main-title div[data-testid="stMarkdownContainer"] p,
#bank-main-title div[data-testid="stMarkdownContainer"] span,
#bank-main-title p, #bank-main-title span {{
    font-size: 2.8rem !important;
    font-family: 'Rye', serif !important;
    font-weight: 900 !important;
    color: #D4A843 !important;
    letter-spacing: 3px !important;
    line-height: 1.2 !important;
}}

/* ẨN SIDEBAR HOÀN TOÀN */
[data-testid="stSidebar"],
[data-testid="stSidebarNav"],
[data-testid="collapsedControl"],
section[data-testid="stSidebar"],
div[data-testid="stSidebarNav"] {{
    display: none !important;
    width: 0 !important;
    min-width: 0 !important;
    visibility: hidden !important;
}}

/* GIỮ NGUYÊN font câu hỏi và đáp án - DÙNG ARIAL */
.bank-question-text,
.bank-question-text * {{
    font-family: 'Arial', 'Helvetica', sans-serif !important;
}}
.bank-answer-text,
.bank-answer-text * {{
    font-family: 'Arial', 'Helvetica', sans-serif !important;
}}
.stRadio label,
.stRadio label span,
.stRadio label p,
.stRadio label div {{
    font-family: 'Arial', 'Helvetica', sans-serif !important;
}}
.paragraph-content-box,
.paragraph-content-box * {{
    font-family: 'Arial', 'Helvetica', sans-serif !important;
}}

/* STYLE NÚT ACTION - NHỎ GỌN, CÙNG HÀNG */
.stButton>button {{
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
    color: #FFFFE0 !important;
    border-radius: 6px !important;
    font-size: 0.72em !important;
    font-weight: 700 !important;
    font-family: 'Rye', serif !important;
    border: 1px solid rgba(255, 255, 255, 0.3) !important;
    padding: 6px 8px !important;
    width: auto !important;
    white-space: normal !important;
    word-break: break-word !important;
    letter-spacing: 0px !important;
    box-shadow: 0 3px 10px rgba(102, 126, 234, 0.4) !important;
    transition: all 0.3s ease !important;
    text-transform: uppercase !important;
}}

.stButton>button:hover {{
    background: linear-gradient(135deg, #764ba2 0%, #667eea 100%) !important;
    box-shadow: 0 8px 25px rgba(118, 75, 162, 0.6) !important;
    transform: translateY(-2px) !important;
    border-color: rgba(255, 255, 255, 0.5) !important;
}}

.stButton>button:active {{
    transform: translateY(0) !important;
    box-shadow: 0 4px 15px rgba(102, 126, 234, 0.3) !important;
}}

/* STYLE CHO NÚT DỊCH (st.toggle) */
div[data-testid="stCheckbox"] label p,
div[data-testid="stCheckbox"] label span,
div[data-testid="stCheckbox"] label div,
div[data-testid="stCheckbox"] label,
div[data-testid="stCheckbox"] p,
div[data-testid="stCheckbox"] span,
div[data-testid="stCheckbox"] div,
.stCheckbox label p,
.stCheckbox label span, 
.stCheckbox label,
.stCheckbox p,
.stCheckbox span {{
    color: #D4A843 !important;
    font-size: 20px !important;
    font-weight: 700 !important;
}}

.stToggle label p {{
    font-size: 20px !important;
    font-weight: 700 !important;
    padding: 0;
    margin: 0;
    line-height: 1 !important;
    color: #D4A843 !important;
}}
.stToggle label,
.stToggle label span,
.stToggle label div,
.stToggle label > div[data-testid="stMarkdownContainer"],
.stToggle label > div[data-testid="stMarkdownContainer"] p,
.stToggle label > div[data-testid="stMarkdownContainer"] span,
.stToggle label * {{
    color: #D4A843 !important;
    font-size: 20px !important;
    font-weight: 700 !important;
}}
.stToggle > label > div[data-testid="stMarkdownContainer"] {{
    margin-top: 10px !important; 
}}

/* Force màu vàng cho toggle text */
[data-testid="stMarkdownContainer"] > p {{
    color: inherit !important;
}}
.stToggle [data-testid="stMarkdownContainer"] > p {{
    color: #D4A843 !important;
}}

/* SELECTBOX - LABEL TIÊU ĐỀ */
div.stSelectbox label p,
div.stSelectbox label span,
div.stSelectbox label,
div[data-testid*="column"] label p {{
    color: #D4A843 !important;
    font-size: 1.35rem !important;
    font-family: 'Rye', serif !important;
    font-weight: bold;
    text-shadow: none;
}}

/* SELECTBOX - KHUNG CHỌN */
.stSelectbox div[data-baseweb="select"] {{
    background-color: rgba(0, 0, 0, 0.7) !important;
    border: 1px solid #C29954 !important;
    border-radius: 8px !important;
}}

/* SELECTBOX - NỘI DUNG ĐANG HIỂN THỊ (giá trị đã chọn) */
.stSelectbox div[data-baseweb="select"] *,
.stSelectbox div[data-baseweb="select"] span,
.stSelectbox div[data-baseweb="select"] div,
.stSelectbox div[data-baseweb="select"] input {{
    font-family: 'Rye', serif !important;
    color: #000000 !important;
    font-size: 1.1rem !important;
    background-color: #FFFFFF !important;
}}

/* SELECTBOX - DROPDOWN LIST OPTIONS (nền sáng, chữ đen) */
[data-baseweb="popover"],
[data-baseweb="popover"] *,
[data-baseweb="menu"],
[data-baseweb="menu"] *,
[role="listbox"],
[role="listbox"] *,
[role="option"],
[role="option"] *,
li[role="option"],
li[role="option"] * {{
    font-family: 'Rye', serif !important;
    font-size: 1.1rem !important;
    color: #000000 !important;
    background-color: #FFFFFF !important;
}}

/* Hover option */
[role="option"]:hover,
[role="option"]:hover * {{
    background-color: rgba(184, 134, 11, 0.3) !important;
    color: #000000 !important;
}}

/* STYLE CHO KHUNG DỊCH - ÁP DỤNG CHO CẢ PC & MOBILE */
div[data-testid="stAlert"] {{
    background-color: rgba(30, 30, 30, 0.95) !important;
    border-left: 4px solid #00d4ff !important;
    border-radius: 8px !important;
    box-shadow: 0 4px 15px rgba(0, 212, 255, 0.3) !important;
}}

/* ============================================================
   GLOBAL: Tất cả text UI (trừ câu hỏi/đáp án) → #FFFFE0
   ============================================================ */

/* Streamlit general text → Arial làm mặc định toàn trang */
.stApp p, .stApp span, .stApp div,
.stApp label,
[data-testid="stText"],
[data-testid="stMarkdownContainer"] *,
[data-testid="stWidgetLabel"] *,
[data-testid="stWidgetLabel"],
div[data-testid="stNotification"] *,
.stWarning *, .stInfo *, .stError *,
.stSuccess * {{
    color: #D4A843 !important;
    font-family: 'Arial', 'Helvetica', sans-serif !important;
}}

/* Chỉ các label UI (selectbox, toggle, button) mới dùng Rye */
div.stSelectbox label, div.stSelectbox label p, div.stSelectbox label span,
.stToggle label, .stToggle label p, .stToggle label span,
.stButton>button {{
    font-family: 'Rye', serif !important;
}}

/* Ngoại lệ: tiêu đề chính giữ màu riêng */
#bank-main-title {{
    margin-top: -120px !important;
    margin-bottom: 10px !important;
    text-align: center !important;
    font-size: 2.8rem !important;
    font-family: 'Rye', serif !important;
    font-weight: 900 !important;
    color: #D4A843 !important;
    letter-spacing: 3px !important;
}}
#bank-main-title span {{
    color: #D4A843 !important;
    font-size: 2.8rem !important;
    font-family: 'Rye', serif !important;
    font-weight: 900 !important;
    letter-spacing: 3px !important;
}}

/* Ngoại lệ: câu hỏi và đáp án giữ màu riêng */
.bank-question-text, .bank-question-text * {{ color: #FF8C00 !important; }}
.bank-answer-text {{ color: #FFFFFF; }}
.stRadio label, .stRadio label span,
.stRadio label p, .stRadio label div {{ color: #FFFFFF !important; }}
.paragraph-content-box, .paragraph-content-box * {{ color: #F0F0F0 !important; }}

/* Ngoại lệ: đáp án đúng/sai vẫn giữ màu xanh/đỏ */
div[data-testid="stAlert"] *,
div[data-testid="stAlert"] p,
div[data-testid="stAlert"] strong,
div[data-testid="stAlert"] em,
div[data-testid="stAlert"] li,
div[data-testid="stAlert"] span,
div[data-testid="stAlert"] div {{ color: #FFFFFF !important; }}
div[data-testid="stAlert"] strong {{ color: #FFD700 !important; }}

/* ĐÁP ÁN SAU KHI NỘP BÀI: CLASS-BASED (KHÔNG BỊ OVERRIDE BỞI GLOBAL RULE) */
/* Đáp án đúng → màu xanh lá */
.bank-answer-text.answer-correct,
div.bank-answer-text.answer-correct,
.stApp .bank-answer-text.answer-correct,
.stApp div.bank-answer-text.answer-correct {{
    color: #00ff00 !important;
    font-weight: 900 !important;
}}
/* Người dùng chọn sai → màu vàng */
.bank-answer-text.answer-selected-wrong,
div.bank-answer-text.answer-selected-wrong,
.stApp .bank-answer-text.answer-selected-wrong,
.stApp div.bank-answer-text.answer-selected-wrong {{
    color: #FFD700 !important;
    font-weight: 900 !important;
}}
/* Đáp án không được chọn và sai → màu trắng */
.bank-answer-text.answer-wrong,
div.bank-answer-text.answer-wrong,
.stApp .bank-answer-text.answer-wrong,
.stApp div.bank-answer-text.answer-wrong {{
    color: #FFFFFF !important;
}}

/* MOBILE RESPONSIVE */
@media (max-width: 768px) {{
    #main-title-container {{ height: 100px; padding-top: 10px; }}
    #main-title-container h1 {{ font-size: 8vw; line-height: 1.5 !important; }}
    .main > div:first-child {{ padding-top: 90px !important; }}

    /* Thu nhỏ tiêu đề chính vừa 1 dòng trên mobile */
    #bank-main-title,
    #bank-main-title span,
    #bank-main-title p,
    #bank-main-title div[data-testid="stMarkdownContainer"] p,
    #bank-main-title div[data-testid="stMarkdownContainer"] span {{
        font-size: 5.2vw !important;
        letter-spacing: 1px !important;
        white-space: nowrap !important;
        overflow: hidden !important;
        text-overflow: ellipsis !important;
    }}

    /* Thu nhỏ label hộp chọn trên mobile */
    div.stSelectbox label p,
    div.stSelectbox label span,
    div.stSelectbox label,
    div[data-testid*="column"] label p {{
        font-size: 1rem !important;
    }}
    
    /* Chỉnh kích thước tiêu đề trên mobile */
    .result-title h3 {{
        font-size: 1.3rem !important;
        font-family: 'Arial', 'Helvetica', sans-serif !important;
        color: #D4A843 !important;
        white-space: normal !important;
        overflow: visible !important;
        text-overflow: clip !important;
        padding: 0 10px !important;
        line-height: 1.3 !important;
    }}
    
    /* Màu vàng cho câu hỏi trên mobile */
    .bank-question-text {{
        color: #FFFF00 !important;
        background-color: rgba(0, 0, 0, 0.75) !important;
        display: inline-block !important; /* BAO VỪA CHỮ */
    }}
    
    /* Nút trên mobile */
    .stButton>button {{
        font-size: 0.64em !important;
        padding: 6px 6px !important;
        border-radius: 6px !important;
        white-space: normal !important;
        word-break: break-word !important;
        letter-spacing: 0px !important;
        border: 1px solid rgba(255, 255, 255, 0.3) !important;
    }}
    
    /* Cập nhật mobile cho đoạn văn */
    .paragraph-title {{
        font-size: 1.2rem;
        padding: 5px 10px;
        margin-top: 10px;
    }}
    .paragraph-content-box {{
        font-size: 16px !important; 
        line-height: 1.4;
        padding: 10px;
    }}
}}
</style>
"""

st.markdown(css_style, unsafe_allow_html=True)

# ====================================================
# 🧭 HEADER & BODY
# ====================================================
st.markdown(f"""
<div id="logo-container">
    <div id="logo-wrap">
        <img src="data:image/jpeg;base64,{img_logo_base64}" alt="Logo" />
    </div>
</div>

<div id="logo2-container">
    <div class="logo2-wrap" id="logo2-wrap">
        <img src="data:image/png;base64,{img_logo2_base64}" alt="Logo2" id="logo2-img"/>
        <svg class="ellipse-border" viewBox="0 0 228 100">
            <style>
                .el-tail {{
                    stroke-dasharray: 100 900;
                    animation: elip-run 2s linear infinite;
                }}
                .el-mid {{
                    stroke-dasharray: 60 940;
                    animation: elip-run 2s linear infinite;
                }}
                .el-tip {{
                    stroke-dasharray: 18 982;
                    animation: elip-run 2s linear infinite;
                }}
                @keyframes elip-run {{
                    from {{ stroke-dashoffset: 1000; }}
                    to   {{ stroke-dashoffset: 0; }}
                }}
            </style>
            <ellipse cx="114" cy="50" rx="112" ry="48"
                fill="none" stroke="#b8860b" stroke-width="2.5"
                stroke-linecap="round" pathLength="1000"
                class="el-tail"/>
            <ellipse cx="114" cy="50" rx="112" ry="48"
                fill="none" stroke="#FFD700" stroke-width="3"
                stroke-linecap="round" pathLength="1000"
                class="el-mid"/>
            <ellipse cx="114" cy="50" rx="112" ry="48"
                fill="none" stroke="#FFF8C0" stroke-width="1.5"
                stroke-linecap="round" pathLength="1000"
                class="el-tip"/>
        </svg>
    </div>
</div>

<div id="header-content-wrapper">
    <div id="main-title-container"></div>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div id="bank-main-title">
    NGÂN HÀNG TRẮC NGHIỆM
</div>
""", unsafe_allow_html=True)

if "current_group_idx" not in st.session_state: st.session_state.current_group_idx = 0
if "submitted" not in st.session_state: st.session_state.submitted = False
if "current_mode" not in st.session_state: st.session_state.current_mode = "group"
if "last_bank_choice" not in st.session_state: st.session_state.last_bank_choice = "----" 
if "doc_selected" not in st.session_state: st.session_state.doc_selected = "Phụ lục 1 : Ngữ pháp chung" 
if 'translations' not in st.session_state: st.session_state.translations = {} # KHỞI TẠO STATE DỊCH THUẬT Q&A
if 'active_translation_key' not in st.session_state: st.session_state.active_translation_key = None # KHỞI TẠO KEY DỊCH Q&A ĐỘC QUYỀN
if 'active_passage_translation' not in st.session_state: st.session_state.active_passage_translation = None # KHỞI TẠO KEY DỊCH ĐOẠN VĂN ĐỘC QUYỀN
if 'passage_translations_cache' not in st.session_state: st.session_state.passage_translations_cache = {} # CACHE DỊCH ĐOẠN VĂN
if 'current_passage_id_displayed' not in st.session_state: st.session_state.current_passage_id_displayed = None 
if 'group_mode_title' not in st.session_state: st.session_state.group_mode_title = "Luyện tập theo nhóm (30 câu/nhóm)"
if 'last_source' not in st.session_state: st.session_state.last_source = ""
if 'group_selector_key' not in st.session_state: st.session_state.group_selector_key = 0
if 'exam_choice_val' not in st.session_state: st.session_state.exam_choice_val = "----"
if 'last_exam_choice' not in st.session_state: st.session_state.last_exam_choice = "----"
if 'bank_choice_val' not in st.session_state: st.session_state.bank_choice_val = "----"

# ============================================================
# BƯỚC 1: CHỌN KỲ THI
# ============================================================
EXAM_OPTIONS = ["----", "Thi Docwise", "Thi CCUQ", "Thi CAAV"]
exam_choice = st.selectbox("Chọn kỳ thi:", EXAM_OPTIONS, index=EXAM_OPTIONS.index(st.session_state.get('exam_choice_val', '----')), key="exam_selector_master")
st.session_state.exam_choice_val = exam_choice

# Reset khi đổi kỳ thi
if st.session_state.get('last_exam_choice') != exam_choice:
    st.session_state.bank_choice_val = "----"
    if 'bank_selector_master' in st.session_state:
        st.session_state.bank_selector_master = "----"
    
    st.session_state.current_group_idx = 0
    st.session_state.submitted = False
    st.session_state.current_mode = "group"
    st.session_state.active_translation_key = None
    st.session_state.active_passage_translation = None
    st.session_state.current_passage_id_displayed = None
    
    last_exam = st.session_state.get('last_exam_choice', '----')
    if isinstance(last_exam, str) and last_exam != "----":
        exam_slug_old = last_exam.split()[-1].lower()
        st.session_state.pop(f"test_{exam_slug_old}_started", None)
        st.session_state.pop(f"test_{exam_slug_old}_submitted", None)
        st.session_state.pop(f"test_{exam_slug_old}_questions", None)
    
    st.session_state.last_exam_choice = exam_choice
    st.rerun()

# ============================================================
# BƯỚC 2: CHỌN NGÂN HÀNG (tuỳ theo kỳ thi)
# ============================================================
bank_choice = "----"
if exam_choice == "Thi CCUQ":
    BANK_OPTIONS_CCUQ = ["----", "Ngân hàng Luật VAECO", "Ngân hàng Kỹ thuật"]
    bank_choice = st.selectbox("Chọn ngân hàng:", BANK_OPTIONS_CCUQ, index=BANK_OPTIONS_CCUQ.index(st.session_state.get('bank_choice_val', '----')), key="bank_selector_master")
    st.session_state.bank_choice_val = bank_choice
elif exam_choice == "Thi CAAV":
    BANK_OPTIONS_CAAV = ["----", "Ngân hàng CAAV Cabin", "Ngân hàng CAAV Law", "Ngân hàng Human Factor"]
    bank_choice = st.selectbox("Chọn ngân hàng:", BANK_OPTIONS_CAAV, index=BANK_OPTIONS_CAAV.index(st.session_state.get('bank_choice_val', '----') if st.session_state.get('bank_choice_val', '----') in BANK_OPTIONS_CAAV else '----'), key="bank_selector_master")
    st.session_state.bank_choice_val = bank_choice

    # Hộp thoại chọn module cho CAAV Law
    if bank_choice == "Ngân hàng CAAV Law":
        if 'caav_law_module' not in st.session_state:
            st.session_state.caav_law_module = None

        CAAV_LAW_MODULE_OPTIONS = ["----", "Module 10.1", "Module 10.2"]
        caav_law_module_choice = st.selectbox(
            "Chọn module:",
            CAAV_LAW_MODULE_OPTIONS,
            index=CAAV_LAW_MODULE_OPTIONS.index(st.session_state.caav_law_module) if st.session_state.caav_law_module in CAAV_LAW_MODULE_OPTIONS else 0,
            key="caav_law_module_selector"
        )
        # Reset khi đổi module
        if st.session_state.caav_law_module != caav_law_module_choice:
            st.session_state.caav_law_module = caav_law_module_choice
            st.session_state.current_group_idx = 0
            st.session_state.submitted = False
            st.session_state.current_mode = "group"
            st.session_state.active_translation_key = None
            st.session_state.active_passage_translation = None
            st.session_state.current_passage_id_displayed = None
            st.rerun()
    else:
        st.session_state.caav_law_module = None
elif exam_choice == "Thi Docwise":
    bank_choice = "Ngân hàng Docwise"  # tự động, không cần chọn
    st.session_state.bank_choice_val = bank_choice

# Reset khi đổi ngân hàng (trong cùng kỳ thi)
if st.session_state.get('last_bank_choice') != bank_choice and bank_choice != "----":
    st.session_state.current_group_idx = 0
    st.session_state.submitted = False
    st.session_state.current_mode = "group"
    st.session_state.active_translation_key = None
    st.session_state.active_passage_translation = None
    st.session_state.current_passage_id_displayed = None
    last_bank_name = st.session_state.get('last_bank_choice')
    if not isinstance(last_bank_name, str) or last_bank_name == "----": last_bank_name = "null bank"
    bank_slug_old = last_bank_name.split()[-1].lower()
    st.session_state.pop(f"test_{bank_slug_old}_started", None)
    st.session_state.pop(f"test_{bank_slug_old}_submitted", None)
    st.session_state.pop(f"test_{bank_slug_old}_questions", None)
    st.session_state.last_bank_choice = bank_choice
    st.rerun()

if exam_choice != "----" and bank_choice != "----":
    # XỬ LÝ LOGIC NGUỒN DỮ LIỆU
    source = ""
    is_docwise = False

    if "Kỹ thuật" in bank_choice:
        source = "cabbank.docx"
    elif "Luật VAECO" in bank_choice:
        source = "lawbank.docx"
    elif "CAAV Cabin" in bank_choice:
        source = "caav cab.docx"
    elif "CAAV Law" in bank_choice:
        caav_law_module = st.session_state.get('caav_law_module', None)
        if caav_law_module == "Module 10.1":
            source = "caav law1.docx"
        elif caav_law_module == "Module 10.2":
            source = "caav law2.docx"
        else:
            st.info("Vui lòng chọn module để tiếp tục.")
            st.stop()
    elif "Human Factor" in bank_choice:
        source = "caav hf.docx"
    elif "Docwise" in bank_choice:
        is_docwise = True
        # Chọn Phụ lục
        doc_options = ["Phụ lục 1 : Ngữ pháp chung", "Phụ lục 2 : Từ vựng, thuật ngữ", "Phụ lục 3 : Bài đọc hiểu", "Phụ lục 4 : Luật và qui trình", "Phụ lục 5 : Chuyên ngành"]
        doc_selected_new = st.selectbox("Chọn Phụ lục:", doc_options, index=doc_options.index(st.session_state.get('doc_selected', doc_options[0])), key="docwise_selector")

        # Xử lý khi đổi phụ lục (reset mode)
        if st.session_state.doc_selected != doc_selected_new:
            st.session_state.doc_selected = doc_selected_new
            st.session_state.current_group_idx = 0
            st.session_state.group_selector_key += 1
            st.session_state.last_source = ""
            st.session_state.submitted = False
            st.session_state.current_mode = "group"
            st.session_state.active_translation_key = None
            st.session_state.active_passage_translation = None
            st.session_state.current_passage_id_displayed = None
            st.rerun()

        if st.session_state.doc_selected == "Phụ lục 1 : Ngữ pháp chung":
            source = "PL1.docx"
        elif st.session_state.doc_selected == "Phụ lục 2 : Từ vựng, thuật ngữ":
            source = "PL2.docx"
        elif st.session_state.doc_selected == "Phụ lục 3 : Bài đọc hiểu":
            source = "PL3.docx"
        elif st.session_state.doc_selected == "Phụ lục 4 : Luật và qui trình":
            source = "PL4.docx"
        elif st.session_state.doc_selected == "Phụ lục 5 : Chuyên ngành":
            source = "PL5.docx"

    # LOAD CÂU HỎI
    questions = []
    if source:
        if "Kỹ thuật" in bank_choice:
            questions = parse_cabbank(source)
        elif "Luật VAECO" in bank_choice:
            questions = parse_lawbank(source)
        elif "CAAV Cabin" in bank_choice:
            questions = parse_cabbank(source)
        elif "CAAV Law" in bank_choice:
            questions = parse_lawbank(source)
        elif "Human Factor" in bank_choice:
            questions = parse_cabbank(source)
        elif is_docwise:
            if source == "PL1.docx":
                questions = parse_pl1(source) # Sử dụng parser cũ (dùng (*))
            elif source == "PL2.docx":
                questions = parse_pl2(source) # Sử dụng parser mới (dùng (*))
            elif source == "PL3.docx":
                questions = parse_pl3_passage_bank(source) # <-- Dùng parser đã sửa cho PL3
            elif source == "PL4.docx":
                questions = parse_pl4_law_process(source)
            elif source == "PL5.docx":
                questions = parse_pl5_specialized(source)
        
        # XỬ LÝ LỌC SỐ THỨ TỰ CHO CÂU HỎI 1 (CAAV LAW 10.2 & HUMAN FACTOR)
        if (bank_choice == "Ngân hàng CAAV Law" and st.session_state.get('caav_law_module') == "Module 10.2") or \
           (bank_choice == "Ngân hàng Human Factor"):
            for q in questions:
                q["question"] = strip_question_number(q["question"])
    
    if not questions:
        # Cập nhật thông báo lỗi để phù hợp với logic (*) cho cả PL1 và PL2
        st.error(f"❌ Không đọc được câu hỏi nào từ file **{source}**. Vui lòng kiểm tra file và cấu trúc thư mục (đảm bảo file nằm trong thư mục gốc hoặc thư mục 'pages/'), và kiểm tra lại định dạng đáp án đúng (dùng dấu `(*)`).")
        st.stop() 
    
    total = len(questions)

   # === LOGIC NHÓM CÂU HỎI THEO MODE (PL3 & PL4 TÙY CHỈNH) - ĐÃ SỬA THEO YÊU CẦU MỚI ===
    group_size = 30 # Mặc định 30 câu/nhóm
    custom_groups = [] # Chỉ dùng cho PL3 & PL4
    is_passage_grouping = False # Đổi tên từ is_pl3_grouping

    if is_docwise and source in ["PL3.docx", "PL4.docx"]:
        is_passage_grouping = True
        passage_groups = {}
        
        # Nhóm câu hỏi theo tên Paragraph
        for q in questions:
            # group_key: "Paragraph 1 ."
            group_key = q.get('group', 'Không có đoạn văn')
            if group_key not in passage_groups:
                passage_groups[group_key] = []
            
            passage_groups[group_key].append(q)
            
        # ----------------------------------------------------
        # LOGIC MỚI: NHÓM 2 PARAGRAPH THÀNH 1 NHÓM
        # ----------------------------------------------------
        passage_names = list(passage_groups.keys())
        
        # Duyệt qua danh sách tên Paragraph theo bước nhảy 2
        for i in range(0, len(passage_names), 2):
            p1_name = passage_names[i]
            p2_name = passage_names[i+1] if i + 1 < len(passage_names) else None
            
            questions_in_pair = passage_groups[p1_name]
            
            # Xử lý Paragraph thứ 2
            if p2_name:
                questions_in_pair.extend(passage_groups[p2_name])
                
                # Bóc tách số thứ tự khỏi chuỗi "Paragraph X ."
                p1_match = re.search(r'Paragraph\s*(\d+)', p1_name, re.I)
                p2_match = re.search(r'Paragraph\s*(\d+)', p2_name, re.I)
                
                p1_num = p1_match.group(1) if p1_match else p1_name
                p2_num = p2_match.group(1) if p2_match else p2_name
                
                base_group_label = f"Paragraph {p1_num} & {p2_num}"
            else:
                # Xử lý Paragraph lẻ cuối cùng (ví dụ: "Paragraph 11")
                p1_match = re.search(r'Paragraph\s*(\d+)', p1_name, re.I)
                p1_num = p1_match.group(1) if p1_match else p1_name
                base_group_label = f"Paragraph {p1_num}"
            
            # TẠO LABEL CUỐI CÙNG (CHỈ DÙNG TÊN PARAGRAPH)
            final_group_label = base_group_label
            
            if questions_in_pair:
                # Dù có câu hỏi hay không, vẫn dùng base_group_label (ví dụ: "Paragraph 1 & 2")
                pass
            else:
                 # Trường hợp không có câu hỏi nào (chỉ để dự phòng, hiếm xảy ra)
                final_group_label = base_group_label

            custom_groups.append({
                'label': final_group_label,
                'questions': questions_in_pair
            })
        
        groups = [g['label'] for g in custom_groups]
        # Khi mới chuyển sang PL3/PL4: mặc định Paragraph 1&2 (index 0)
        if st.session_state.get('last_source', '') != source:
            st.session_state.current_group_idx = 0
            st.session_state.last_source = source
        elif st.session_state.current_group_idx >= len(groups):
            st.session_state.current_group_idx = 0
    else:
        # Nhóm câu hỏi theo số lượng (30 câu/nhóm) cho các ngân hàng khác
        groups = [f"Câu {i*group_size+1}-{min((i+1)*group_size, total)}" for i in range(math.ceil(total/group_size))]
        if st.session_state.get('last_source', '') != source:
            st.session_state.current_group_idx = 0
            st.session_state.last_source = source
        
    # --- MODE: GROUP ---
    if st.session_state.current_mode == "group":
        if total > 0:
            if st.session_state.current_group_idx >= len(groups): st.session_state.current_group_idx = 0
            selected = st.selectbox(
                "Chọn nhóm câu:", groups,
                index=st.session_state.current_group_idx,
                key=f"group_selector_{st.session_state.group_selector_key}"
            )
            
            # Xử lý khi chuyển nhóm câu
            new_idx = groups.index(selected)
            if st.session_state.current_group_idx != new_idx:
                st.session_state.current_group_idx = new_idx
                st.session_state.submitted = False
                st.session_state.active_translation_key = None # Reset dịch Q&A
                st.session_state.active_passage_translation = None # Reset dịch Passage
                st.session_state.current_passage_id_displayed = None # Reset passage display
                st.rerun()

            idx = st.session_state.current_group_idx
            
            if is_passage_grouping:
                batch = custom_groups[idx]['questions']
                start = 0 # Not relevant in this new grouping mode
            else:
                # Logic lấy batch cũ (30 câu/nhóm)
                start = idx * group_size
                end = min((idx+1) * group_size, total)
                batch = questions[start:end]

            # Set starting index for questions in non-PL3 mode
            start_i = start + 1 
            
            st.markdown('<div style="margin-top: 12px;"></div>', unsafe_allow_html=True)
            
            # Xác định nhãn nút hiển thị toàn bộ
            if is_docwise:
                pl_short = st.session_state.get('doc_selected', '').split(':')[0].strip()
                btn_all_label = f"📖 Hiển thị {pl_short}"
            else:
                btn_all_label = "📖 Hiển thị toàn bộ ngân hàng"

            # Hàng 1: nút Hiển thị — full width, font nhỏ để vừa 1 dòng
            if st.button(btn_all_label, key="btn_show_all", use_container_width=True):
                st.session_state.current_mode = "all"
                st.session_state.all_submitted = False
                st.session_state.active_translation_key = None
                st.session_state.active_passage_translation = None
                st.session_state.current_passage_id_displayed = None
                st.rerun()

            # Hàng 2: nút Làm bài test — canh giữa bằng cột đối xứng
            _, col_test, _ = st.columns([1, 1, 1])
            with col_test:
                if is_docwise:
                    pl_short = st.session_state.get('doc_selected', '').split(':')[0].strip()
                    if st.button(f"📝 Làm bài test {pl_short}", key="btn_start_appendix_test", use_container_width=True):
                        st.session_state.current_mode = "appendix_test"
                        st.session_state.active_translation_key = None
                        st.session_state.active_passage_translation = None
                        st.session_state.current_passage_id_displayed = None
                        st.rerun()
                else:
                    if st.button("📝 Làm bài test", key="btn_start_test", use_container_width=True):
                        st.session_state.current_mode = "test"
                        st.session_state.active_translation_key = None
                        st.session_state.active_passage_translation = None
                        st.session_state.current_passage_id_displayed = None
                        bank_slug_new = bank_choice.split()[-1].lower()
                        test_key_prefix = f"test_{bank_slug_new}"
                        st.session_state.pop(f"{test_key_prefix}_started", None)
                        st.session_state.pop(f"{test_key_prefix}_submitted", None)
                        st.session_state.pop(f"{test_key_prefix}_questions", None)
                        st.rerun()

            # Hàng 3: nút Thi thử Docwise (chỉ khi là Docwise) hoặc Thi thử CAAV (chỉ khi là CAAV)
            if is_docwise:
                _, col_docwise, _ = st.columns([1, 1, 1])
                with col_docwise:
                    if st.button("🎓 Thi thử Docwise", key="btn_start_docwise_test", use_container_width=True):
                        st.session_state.current_mode = "test"
                        st.session_state.active_translation_key = None
                        st.session_state.active_passage_translation = None
                        st.session_state.current_passage_id_displayed = None
                        st.rerun()
            elif exam_choice == "Thi CAAV":
                _, col_caav_mock, _ = st.columns([1, 1, 1])
                with col_caav_mock:
                    if st.button("🎓 Thi thử CAAV", key="btn_start_caav_mock", use_container_width=True):
                        st.session_state.current_mode = "caav_mock_test"
                        st.session_state.active_translation_key = None
                        st.session_state.active_passage_translation = None
                        st.session_state.current_passage_id_displayed = None
                        st.rerun()
            st.markdown('<div class="question-separator"></div>', unsafe_allow_html=True)
            
            
            # --- BẮT ĐẦU VÒNG LẶP CÂU HỎI ---
            if batch:
                current_passage_id_in_group_mode = None
                
                if not st.session_state.submitted:
                    # Luyện tập
                    for i_local, q in enumerate(batch):
                        i_global = q.get('global_number', start + i_local + 1) # Sử dụng global_number nếu có
                        q_key = f"q_{i_global}_{hash(q['question'])}" 
                        translation_key = f"trans_{q_key}"
                        is_active = (translation_key == st.session_state.active_translation_key)
                        
                        # --- CẬP NHẬT: HIỂN THỊ ĐOẠN VĂN (CHO PL3) TRƯỚC CÂU HỎI ---
                        passage_content = q.get('paragraph_content', '').strip()
                        group_name = q.get('group', '')
                        
                        if passage_content:
                            passage_id = f"passage_{group_name}_{hash(passage_content)}"
                            is_passage_active = (passage_id == st.session_state.active_passage_translation)

                            if passage_id != current_passage_id_in_group_mode:
                                # 1. In đậm, đổi màu tiêu đề
                                st.markdown(f'<div class="paragraph-title">**{group_name}**</div>', unsafe_allow_html=True) 
                                
                                # 2. Hiển thị nội dung đoạn văn gốc
                                st.markdown(f'<div class="paragraph-content-box">{passage_content}</div>', unsafe_allow_html=True)
                                
                                # 3. Thêm Nút Dịch Đoạn Văn
                                if not q.get('is_content_only'):
                                    st.toggle(
                                        "🌐 Dịch đoạn văn sang Tiếng Việt", 
                                        value=is_passage_active, 
                                        key=f"toggle_passage_{passage_id}",
                                        on_change=on_passage_translate_toggle,
                                        args=(passage_id,)
                                    )
                                
                                # 4. Hiển thị Bản Dịch Đoạn Văn
                                if is_passage_active and not q.get('is_content_only'):
                                    translated_passage = st.session_state.passage_translations_cache.get(passage_id)
                                    if not isinstance(translated_passage, str):
                                        # GỌI HÀM DỊCH CHỈ ĐOẠN VĂN
                                        translated_passage = translate_passage_content(passage_content)
                                        st.session_state.passage_translations_cache[passage_id] = translated_passage

                                    st.markdown(f"""
                                    <div data-testid="stAlert" class="stAlert stAlert-info">
                                        <div style="font-size: 18px; line-height: 1.6; color: white; padding: 10px;">
                                            <strong style="color: #FFD700;">[Bản dịch Đoạn văn]</strong>
                                            <div class="paragraph-content-box" style="white-space: pre-wrap; margin-bottom: 0px; padding: 10px; background-color: rgba(0, 0, 0, 0.5); border-left: 3px solid #00d4ff;">
                                            {translated_passage}
                                            </div>
                                        </div>
                                    </div>
                                    """, unsafe_allow_html=True)

                                st.markdown("---") 
                                
                                current_passage_id_in_group_mode = passage_id
                        # -----------------------------------------------------------------
                        
                        if q.get('is_content_only'):
                            continue
                        
                        # Fix số trùng: chỉ dùng số cục bộ khi có đoạn văn thực sự (PL3/PL4)
                        if q.get('group', '').startswith('Paragraph') and q.get('paragraph_content'):
                            display_num = q.get('number', i_global)
                        else:
                            display_num = i_global
                        
                        # Hiển thị câu hỏi
                        st.markdown(f'<div class="bank-question-text">{display_num}. {q["question"]}</div>', unsafe_allow_html=True) 

                        # Nút Dịch Q&A ở dưới
                        st.toggle(
                            "🌐 Dịch Câu hỏi & Đáp án sang Tiếng Việt", 
                            value=is_active, 
                            key=f"toggle_{translation_key}",
                            on_change=on_translate_toggle,
                            args=(translation_key,)
                        )

                        # Hiển thị Bản Dịch Q&A
                        if is_active:
                            # Check if translated content is already cached
                            translated_content = st.session_state.translations.get(translation_key)
                            
                            # If not cached or is not a string (default True/False state)
                            if not isinstance(translated_content, str):
                                # GỌI HÀM MỚI ĐỂ GỬI CHỈ CÂU HỎI VÀ ĐÁP ÁN ĐI DỊCH
                                full_text_to_translate = build_translation_text_for_qa(q) 
                                st.session_state.translations[translation_key] = translate_text(full_text_to_translate)
                                translated_content = st.session_state.translations[translation_key]

                            st.info(translated_content, icon="🌐")

                        # Hiển thị Radio Button
                        default_val = st.session_state.get(q_key, None)
                        # Ẩn dấu (*) khỏi nhãn của radio button
                        options_clean = [opt.replace("(*)", "").strip() for opt in q["options"]]
                        st.radio("", options_clean, index=options_clean.index(default_val.replace("(*)", "").strip()) if default_val and default_val.replace("(*)", "").strip() in options_clean else None, key=q_key)
                        st.markdown('<div class="question-separator"></div>', unsafe_allow_html=True)
                    if st.button("✅ Nộp bài", key="submit_group"):
                        st.session_state.submitted = True
                        st.session_state.scroll_to_group_result = True
                        st.session_state.active_translation_key = None # Tắt dịch Q&A khi nộp
                        st.session_state.active_passage_translation = None # Tắt dịch Passage khi nộp
                        st.rerun()
                else:
                    # Chế độ xem đáp án
                    score = 0
                    for i_local, q in enumerate(batch):
                        i_global = q.get('global_number', start + i_local + 1)
                        q_key = f"q_{i_global}_{hash(q['question'])}" 
                        selected_opt = st.session_state.get(q_key)
                        correct = clean_text(q["answer"])
                        is_correct = (selected_opt is not None) and (clean_text(selected_opt) == correct)
                        translation_key = f"trans_{q_key}"
                        is_active = (translation_key == st.session_state.active_translation_key)
                        
                        # --- CẬP NHẬT: HIỂN THỊ ĐOẠN VĂN (CHO PL3) TRƯỚC CÂU HỎI ---
                        passage_content = q.get('paragraph_content', '').strip()
                        group_name = q.get('group', '')
                        
                        if passage_content:
                            passage_id = f"passage_{group_name}_{hash(passage_content)}"
                            is_passage_active = (passage_id == st.session_state.active_passage_translation)

                            if passage_id != current_passage_id_in_group_mode:
                                # 1. In đậm, đổi màu tiêu đề
                                st.markdown(f'<div class="paragraph-title">**{group_name}**</div>', unsafe_allow_html=True) 
                                
                                # 2. Hiển thị nội dung đoạn văn gốc
                                st.markdown(f'<div class="paragraph-content-box">{passage_content}</div>', unsafe_allow_html=True)
                                
                                # 3. Thêm Nút Dịch Đoạn Văn
                                if not q.get('is_content_only'):
                                    st.toggle(
                                        "🌐 Dịch đoạn văn sang Tiếng Việt", 
                                        value=is_passage_active, 
                                        key=f"toggle_passage_{passage_id}",
                                        on_change=on_passage_translate_toggle,
                                        args=(passage_id,)
                                    )
                                
                                # 4. Hiển thị Bản Dịch Đoạn Văn
                                if is_passage_active and not q.get('is_content_only'):
                                    translated_passage = st.session_state.passage_translations_cache.get(passage_id)
                                    if not isinstance(translated_passage, str):
                                        # GỌI HÀM DỊCH CHỈ ĐOẠN VĂN
                                        translated_passage = translate_passage_content(passage_content)
                                        st.session_state.passage_translations_cache[passage_id] = translated_passage

                                    st.markdown(f"""
                                    <div data-testid="stAlert" class="stAlert stAlert-info">
                                        <div style="font-size: 18px; line-height: 1.6; color: white; padding: 10px;">
                                            <strong style="color: #FFD700;">[Bản dịch Đoạn văn]</strong>
                                            <div class="paragraph-content-box" style="white-space: pre-wrap; margin-bottom: 0px; padding: 10px; background-color: rgba(0, 0, 0, 0.5); border-left: 3px solid #00d4ff;">
                                            {translated_passage}
                                            </div>
                                        </div>
                                    </div>
                                    """, unsafe_allow_html=True)

                                st.markdown("---") 
                                
                                current_passage_id_in_group_mode = passage_id
                        # -----------------------------------------------------------------
                        if q.get('is_content_only'):
                            continue

                        # Hiển thị câu hỏi: FIX số trùng cho CAAV/LAWBANK
                        if q.get('group', '').startswith('Paragraph') and q.get('paragraph_content'):
                            # Dùng số thứ tự cục bộ (number) chỉ khi là bài đọc hiểu có đoạn văn
                            display_num = q.get('number', i_global) 
                        else:
                            # Dùng số thứ tự toàn cục (i_global) cho tất cả ngân hàng khác
                            display_num = i_global 
                        st.markdown(f'<div class="bank-question-text">{display_num}. {q["question"]}</div>', unsafe_allow_html=True) 

                        # Nút Dịch Q&A ở dưới
                        st.toggle(
                            "🌐 Dịch Câu hỏi & Đáp án sang Tiếng Việt", 
                            value=is_active, 
                            key=f"toggle_{translation_key}",
                            on_change=on_translate_toggle,
                            args=(translation_key,)
                        )

                        # Hiển thị Bản Dịch Q&A
                        if is_active:
                            # Check if translated content is already cached
                            translated_content = st.session_state.translations.get(translation_key)
                            
                            # If not cached or is not a string (default True/False state)
                            if not isinstance(translated_content, str):
                                # GỌI HÀM MỚI ĐỂ GỬI CHỈ CÂU HỎI VÀ ĐÁP ÁN ĐI DỊCH
                                full_text_to_translate = build_translation_text_for_qa(q)
                                st.session_state.translations[translation_key] = translate_text(full_text_to_translate)
                                translated_content = st.session_state.translations[translation_key]

                            st.info(translated_content, icon="🌐")

                        # Hiển thị Đáp án (KẾT QUẢ)
                        for opt in q["options"]:
                            # Ẩn dấu (*) nếu chưa nộp bài
                            opt_display = opt.replace("(*)", "").strip()
                            opt_clean = clean_text(opt)
                            if opt_clean == correct:
                                ans_class = "bank-answer-text answer-correct"
                            elif selected_opt is not None and clean_text(selected_opt) == opt_clean:
                                ans_class = "bank-answer-text answer-selected-wrong"
                            else:
                                ans_class = "bank-answer-text answer-wrong"
                            st.markdown(f'<div class="{ans_class}">{opt_display}</div>', unsafe_allow_html=True)
                        
                        if is_correct: 
                            score += 1
                        st.markdown('<div class="question-separator"></div>', unsafe_allow_html=True) 

                    st.markdown('<div id="group-result"></div>', unsafe_allow_html=True)
                    if st.session_state.pop("scroll_to_group_result", False):
                        st.markdown("""
                        <script>
setTimeout(function() {
    // Streamlit dùng overflow scroll trên div gốc, không phải window
    var containers = [
        document.querySelector('.main .block-container'),
        document.querySelector('[data-testid="stAppViewContainer"] > section'),
        document.querySelector('.main'),
        document.documentElement,
        document.body
    ];
    for (var c of containers) {
        if (c) { c.scrollTop = c.scrollHeight; }
    }
    window.scrollTo(0, document.body.scrollHeight);
    window.parent.scrollTo(0, window.parent.document.body.scrollHeight);
}, 400);
</script>
                        """, unsafe_allow_html=True)
                    st.markdown(f'<div class="result-title"><h3>🎯 KẾT QUẢ: {score}/{len(batch)}</h3></div>', unsafe_allow_html=True)
                    col_reset, col_next = st.columns(2)
                    with col_reset:
                        if st.button("🔄 Làm lại nhóm này", key="reset_group"):
                            # Xoá session state của các radio button trong nhóm
                            for i_local, q in enumerate(batch):
                                i_global = q.get('global_number', start + i_local + 1)
                                st.session_state.pop(f"q_{i_global}_{hash(q['question'])}", None) 
                            st.session_state.submitted = False
                            st.session_state.active_translation_key = None # Reset dịch Q&A
                            st.session_state.active_passage_translation = None # Reset dịch Passage
                            st.rerun()
                    with col_next:
                        if st.session_state.current_group_idx < len(groups) - 1:
                            if st.button("➡️ Tiếp tục nhóm sau", key="next_group"):
                                st.session_state.current_group_idx += 1
                                st.session_state.submitted = False
                                st.session_state.active_translation_key = None # Reset dịch Q&A
                                st.session_state.active_passage_translation = None # Reset dịch Passage
                                st.rerun()
                        else: st.info("🎉 Đã hoàn thành tất cả các nhóm câu hỏi!")
            else: st.warning("Không có câu hỏi trong nhóm này.")
        else: st.warning("Không có câu hỏi nào trong ngân hàng này.")

    elif st.session_state.current_mode == "all":
        if st.button("⬅️ Quay lại chế độ Luyện tập theo nhóm"):
            st.session_state.current_mode = "group"
            st.session_state.active_translation_key = None # Reset dịch Q&A
            st.session_state.active_passage_translation = None # Reset dịch Passage
            st.session_state.current_passage_id_displayed = None # Reset passage display
            st.rerun()
        st.markdown('<div class="question-separator"></div>', unsafe_allow_html=True)
        display_all_questions(questions)
        
    elif st.session_state.current_mode == "appendix_test":
        if st.button("⬅️ Quay lại chế độ Luyện tập theo nhóm"):
            st.session_state.current_mode = "group"
            st.session_state.active_translation_key = None
            st.session_state.active_passage_translation = None
            st.session_state.current_passage_id_displayed = None
            st.rerun()
        st.markdown('<div class="question-separator"></div>', unsafe_allow_html=True)
        display_appendix_test_mode(st.session_state.doc_selected)

    elif st.session_state.current_mode == "caav_mock_test":
        if st.button("⬅️ Quay lại chế độ Luyện tập theo nhóm"):
            st.session_state.current_mode = "group"
            st.session_state.active_translation_key = None
            st.session_state.active_passage_translation = None
            st.session_state.current_passage_id_displayed = None
            st.rerun()
        st.markdown('<div class="question-separator"></div>', unsafe_allow_html=True)
        display_caav_mock_test_mode()

    elif st.session_state.current_mode == "test":
        if st.button("⬅️ Quay lại chế độ Luyện tập theo nhóm"):
            st.session_state.current_mode = "group"
            st.session_state.active_translation_key = None # Reset dịch Q&A
            st.session_state.active_passage_translation = None # Reset dịch Passage
            st.session_state.current_passage_id_displayed = None # Reset passage display
            st.rerun()
        st.markdown('<div class="question-separator"></div>', unsafe_allow_html=True)
        if is_docwise:
            display_docwise_test_mode(bank_choice)
        else:
            display_test_mode(questions, bank_choice)
